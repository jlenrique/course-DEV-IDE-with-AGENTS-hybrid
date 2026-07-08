"""Deterministic syllabus readers and module-metadata proposals."""

from __future__ import annotations

import re
from email import policy
from email.parser import BytesParser
from html.parser import HTMLParser
from pathlib import Path
from typing import Literal

import yaml
from docx import Document
from pydantic import BaseModel, ConfigDict, Field, model_validator

from app.marcus.course_source.models import GapEntry
from app.marcus.course_source.registry import load_course
from app.marcus.lesson_plan.event_type_registry import OPEN_ID_REGEX_PATTERN

SCHEMA_VERSION = "0.1"


class SyllabusExtractionError(ValueError):
    """Raised when a syllabus cannot be structurally extracted."""


class SourceAnchor(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True, validate_assignment=True)

    path: str = Field(min_length=1)
    locator: str = Field(min_length=1)


class AnchoredText(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True, validate_assignment=True)

    value: str = Field(min_length=1)
    source_ref: SourceAnchor


class AnchoredObjective(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True, validate_assignment=True)

    objective_id: str = Field(pattern=OPEN_ID_REGEX_PATTERN)
    text: str = Field(min_length=1)
    source_ref: SourceAnchor


class ExtractedSyllabusDocument(BaseModel):
    model_config = ConfigDict(extra="forbid", validate_assignment=True)

    source_path: str
    source_format: Literal["docx", "mhtml"]
    paragraphs: tuple[str, ...] = Field(default_factory=tuple)
    tables: tuple[tuple[tuple[str, ...], ...], ...] = Field(default_factory=tuple)


class ModuleMetadataProposal(BaseModel):
    model_config = ConfigDict(extra="forbid", validate_assignment=True)

    module_id: str = Field(pattern=OPEN_ID_REGEX_PATTERN)
    status: Literal["proposed", "missing"] = "proposed"
    proposed_slug: AnchoredText | None = None
    slug_status: Literal["existing_aligned", "synthesized_requires_review", "missing"]
    slug_review_note: AnchoredText | None = None
    title: AnchoredText | None = None
    topics: tuple[AnchoredText, ...] = Field(default_factory=tuple)
    source_bucket_suggestions: tuple[AnchoredText, ...] = Field(default_factory=tuple)
    source_refs: tuple[SourceAnchor, ...] = Field(default_factory=tuple)

    @model_validator(mode="after")
    def _proposed_records_have_required_anchors(self) -> ModuleMetadataProposal:
        if self.status == "missing":
            if self.proposed_slug is not None or self.title is not None or self.topics:
                raise ValueError("missing module metadata cannot carry proposed fields")
            return self
        if self.proposed_slug is None or self.title is None:
            raise ValueError("proposed module metadata requires slug and title anchors")
        if not self.source_refs:
            raise ValueError("proposed module metadata requires source_refs")
        return self


class CourseMetadataProposal(BaseModel):
    model_config = ConfigDict(extra="forbid", validate_assignment=True)

    schema_version: Literal["0.1"] = SCHEMA_VERSION
    course_id: str = Field(pattern=OPEN_ID_REGEX_PATTERN)
    source_path: str
    extraction_status: Literal["verified", "format_unsupported"]
    course_title: AnchoredText | None = None
    learner_profile: AnchoredText | None = None
    course_learning_objectives: tuple[AnchoredObjective, ...] = Field(default_factory=tuple)
    modules: tuple[ModuleMetadataProposal, ...] = Field(default_factory=tuple)
    gaps: tuple[GapEntry, ...] = Field(default_factory=tuple)


class _HtmlTableParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.paragraphs: list[str] = []
        self.tables: list[list[list[str]]] = []
        self._table: list[list[str]] | None = None
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag == "table":
            self._table = []
        elif tag == "tr" and self._table is not None:
            self._row = []
        elif tag in {"td", "th"} and self._row is not None:
            self._cell = []

    def handle_data(self, data: str) -> None:
        text = _clean(data)
        if not text:
            return
        if self._cell is not None:
            self._cell.append(text)
        elif self._table is None:
            self.paragraphs.append(text)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"td", "th"} and self._cell is not None and self._row is not None:
            self._row.append(_clean(" ".join(self._cell)))
            self._cell = None
        elif tag == "tr" and self._row is not None and self._table is not None:
            if any(cell for cell in self._row):
                self._table.append(self._row)
            self._row = None
        elif tag == "table" and self._table is not None:
            self.tables.append(self._table)
            self._table = None


def _clean(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split()).strip()


def _looks_mojibake(value: str) -> bool:
    lowered = value.lower()
    suspicious_markers = ("\ufffd", "Ã", "Â", "â€™", "â€œ", "â€", "�")
    if any(marker.lower() in lowered for marker in suspicious_markers):
        return True
    if not value:
        return False
    replacement_ratio = value.count("\ufffd") / len(value)
    return replacement_ratio > 0.001


def _anchor(source_path: str, locator: str) -> SourceAnchor:
    return SourceAnchor(path=source_path, locator=locator)


def _anchored(value: str, source_path: str, locator: str) -> AnchoredText:
    return AnchoredText(value=value, source_ref=_anchor(source_path, locator))


def _slugify(value: str) -> str:
    tokens = re.findall(r"[a-z0-9]+", value.lower())
    return "-".join(tokens[:8]) or "metadata-proposal"


def _strip_page_markers(value: str) -> str:
    return _clean(re.sub(r"\bPage:\s*", "", value))


def _split_cell_items(value: str) -> tuple[str, ...]:
    text = _strip_page_markers(value)
    if not text:
        return ()
    pieces = re.split(r"(?<=[.!?])\s+|(?:\s{2,})", text)
    cleaned = tuple(_clean(piece) for piece in pieces if _clean(piece))
    return cleaned or (text,)


def _find_table(
    document: ExtractedSyllabusDocument, *headers: str
) -> tuple[int, tuple[tuple[str, ...], ...]] | None:
    wanted = tuple(header.lower() for header in headers)
    for index, table in enumerate(document.tables):
        if not table:
            continue
        header = " ".join(table[0]).lower()
        if all(item in header for item in wanted):
            return index, table
    return None


def read_docx_syllabus(path: Path) -> ExtractedSyllabusDocument:
    try:
        document = Document(path)
    except Exception as exc:  # pragma: no cover - library-specific subclasses vary
        raise SyllabusExtractionError(f"docx extraction failed: {exc}") from exc
    paragraphs = tuple(
        _clean(paragraph.text) for paragraph in document.paragraphs if _clean(paragraph.text)
    )
    tables = tuple(
        tuple(tuple(_clean(cell.text) for cell in row.cells) for row in table.rows)
        for table in document.tables
    )
    if not paragraphs and not tables:
        raise SyllabusExtractionError("docx extraction produced no text")
    return ExtractedSyllabusDocument(
        source_path=path.as_posix(),
        source_format="docx",
        paragraphs=paragraphs,
        tables=tables,
    )


def read_mhtml_syllabus(path: Path) -> ExtractedSyllabusDocument:
    try:
        message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    except Exception as exc:  # pragma: no cover - parser internals vary
        raise SyllabusExtractionError(f"mhtml extraction failed: {exc}") from exc
    html = None
    for part in message.walk():
        if part.get_content_type() == "text/html":
            html = part.get_content()
            break
    if not isinstance(html, str) or not html.strip():
        raise SyllabusExtractionError("mhtml extraction found no text/html part")
    if _looks_mojibake(html):
        raise SyllabusExtractionError("mhtml decode produced suspicious text")
    parser = _HtmlTableParser()
    parser.feed(html)
    tables = tuple(tuple(tuple(cell for cell in row) for row in table) for table in parser.tables)
    paragraphs = tuple(parser.paragraphs)
    if not paragraphs and not tables:
        raise SyllabusExtractionError("mhtml extraction produced no text")
    return ExtractedSyllabusDocument(
        source_path=path.as_posix(),
        source_format="mhtml",
        paragraphs=paragraphs,
        tables=tables,
    )


def read_syllabus(path: Path) -> ExtractedSyllabusDocument:
    if path.suffix.lower() == ".docx":
        return read_docx_syllabus(path)
    if path.suffix.lower() == ".doc":
        return read_mhtml_syllabus(path)
    raise SyllabusExtractionError(f"unsupported syllabus format: {path.suffix}")


def render_extracted_document_yaml(document: ExtractedSyllabusDocument) -> str:
    return yaml.safe_dump(document.model_dump(mode="json"), sort_keys=False, allow_unicode=True)


def load_extracted_document_yaml(path: Path) -> ExtractedSyllabusDocument:
    return ExtractedSyllabusDocument.model_validate(
        yaml.safe_load(path.read_text(encoding="utf-8"))
    )


def _course_title(document: ExtractedSyllabusDocument) -> AnchoredText | None:
    source = document.source_path
    if document.source_format == "docx" and len(document.paragraphs) > 1:
        return _anchored(document.paragraphs[1], source, "paragraph 2")
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table):
            if len(row) >= 2 and _clean(row[0]).lower().startswith("course name"):
                return _anchored(row[1], source, f"table {table_index} row {row_index}")
    return None


def _learner_profile(document: ExtractedSyllabusDocument) -> AnchoredText | None:
    source = document.source_path
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.lower().startswith("audience:"):
            return _anchored(
                paragraph.removeprefix("Audience:").strip(), source, f"paragraph {index + 1}"
            )
    for table_index, table in enumerate(document.tables):
        if (
            table
            and _clean(table[0][0]).lower().startswith("course description")
            and len(table) > 1
        ):
            return _anchored(table[1][0], source, f"table {table_index} row 1")
    return None


def _docx_objectives(document: ExtractedSyllabusDocument) -> tuple[AnchoredObjective, ...]:
    try:
        start = (
            document.paragraphs.index(
                "Upon successful completion of this course, participants will be able to:"
            )
            + 1
        )
        end = document.paragraphs.index("Weekly Schedule & Topics")
    except ValueError:
        return ()
    return tuple(
        AnchoredObjective(
            objective_id=f"lo-{index:03d}",
            text=text,
            source_ref=_anchor(document.source_path, f"paragraph {start + index}"),
        )
        for index, text in enumerate(document.paragraphs[start:end], start=1)
    )


def _table_objectives(document: ExtractedSyllabusDocument) -> tuple[AnchoredObjective, ...]:
    for table_index, table in enumerate(document.tables):
        objective_rows = [
            row for row in table if len(row) >= 2 and re.fullmatch(r"[A-Z]+-\d{2}", row[0])
        ]
        if objective_rows:
            return tuple(
                AnchoredObjective(
                    objective_id=row[0].lower(),
                    text=row[1],
                    source_ref=_anchor(
                        document.source_path, f"table {table_index} row {row_index}"
                    ),
                )
                for row_index, row in enumerate(objective_rows, start=1)
            )
    return ()


def _course_objectives(document: ExtractedSyllabusDocument) -> tuple[AnchoredObjective, ...]:
    if document.source_format == "docx":
        return _docx_objectives(document)
    return _table_objectives(document)


def _bucket_suggestions(
    header_cells: tuple[str, ...],
    row: tuple[str, ...],
    source_path: str,
    table_index: int,
    row_index: int,
) -> tuple[AnchoredText, ...]:
    buckets: list[AnchoredText] = []
    mappings = {
        "discussion": "discussion-prompts",
        "readings": "readings",
        "resources": "readings",
        "lecture": "lecture",
        "presentation": "lecture",
        "assignments": "assignments",
        "assessment": "assessments",
        "evaluation": "assessments",
        "lab": "lab",
    }
    for column_index, cell in enumerate(row[1:], start=1):
        if not _clean(cell):
            continue
        header = header_cells[column_index].lower() if column_index < len(header_cells) else ""
        for marker, bucket in mappings.items():
            if marker in header and bucket not in {item.value for item in buckets}:
                buckets.append(
                    _anchored(
                        bucket,
                        source_path,
                        f"table {table_index} row {row_index} column {column_index}",
                    )
                )
    return tuple(buckets)


def _docx_modules(
    document: ExtractedSyllabusDocument, module_ids: list[str]
) -> tuple[ModuleMetadataProposal, ...]:
    match = _find_table(document, "week", "theme", "topics")
    if match is None:
        return _missing_modules(document.source_path, module_ids, "weekly schedule table missing")
    table_index, table = match
    header = table[0]
    modules: list[ModuleMetadataProposal] = []
    for index, row in enumerate(table[1:], start=1):
        module_id = module_ids[index - 1] if index <= len(module_ids) else f"module-{index:02d}"
        title = row[1] if len(row) > 1 else ""
        topics_cell = row[2] if len(row) > 2 else ""
        locator = f"table {table_index} row {index}"
        topics = tuple(
            _anchored(item, document.source_path, f"{locator} topics")
            for item in _split_cell_items(topics_cell)
        )
        modules.append(
            ModuleMetadataProposal(
                module_id=module_id,
                proposed_slug=_anchored(module_id, document.source_path, locator),
                slug_status="existing_aligned",
                slug_review_note=_anchored(
                    "Existing module id retained; syllabus row provides aligned "
                    "title/topic evidence.",
                    document.source_path,
                    locator,
                ),
                title=_anchored(title, document.source_path, f"{locator} theme"),
                topics=topics,
                source_bucket_suggestions=_bucket_suggestions(
                    header, row, document.source_path, table_index, index
                ),
                source_refs=(_anchor(document.source_path, locator),),
            )
        )
    if len(modules) < len(module_ids):
        modules.extend(
            _missing_modules(
                document.source_path,
                module_ids[len(modules) :],
                "syllabus row missing for declared module",
            )
        )
    return tuple(modules)


def _mhtml_modules(
    document: ExtractedSyllabusDocument, module_ids: list[str]
) -> tuple[ModuleMetadataProposal, ...]:
    match = _find_table(document, "week", "discussion", "assignments")
    if match is None:
        return _missing_modules(document.source_path, module_ids, "course calendar table missing")
    table_index, table = match
    header = table[0]
    modules: list[ModuleMetadataProposal] = []
    for index, row in enumerate(table[1:], start=1):
        module_id = module_ids[index - 1] if index <= len(module_ids) else f"module-{index:02d}"
        locator = f"table {table_index} row {index}"
        title_column = 1
        title_text = _split_cell_items(row[title_column] if len(row) > title_column else "")
        fallback_column = 4
        fallback_title = _split_cell_items(
            row[fallback_column] if len(row) > fallback_column else ""
        )
        if not title_text and not fallback_title:
            modules.append(
                _missing_modules(
                    document.source_path,
                    [module_id],
                    "syllabus row present but no title-bearing cells were extracted",
                )[0]
            )
            continue
        title = (title_text or fallback_title)[0]
        title_column = title_column if title_text else fallback_column
        combined = " ".join(_strip_page_markers(cell) for cell in row[1:] if _clean(cell))
        proposed_slug = f"module-{index:02d}-{_slugify(combined or title)}"
        topics = tuple(
            _anchored(item, document.source_path, f"{locator} column {column_index}")
            for column_index, cell in enumerate(row[1:], start=1)
            for item in _split_cell_items(cell)
        )
        modules.append(
            ModuleMetadataProposal(
                module_id=module_id,
                proposed_slug=_anchored(proposed_slug, document.source_path, locator),
                slug_status="synthesized_requires_review",
                slug_review_note=_anchored(
                    "Deterministic slug synthesized from the syllabus row; "
                    "operator review is required before any B-2 rename.",
                    document.source_path,
                    locator,
                ),
                title=_anchored(
                    title,
                    document.source_path,
                    f"{locator} column {title_column}",
                ),
                topics=topics,
                source_bucket_suggestions=_bucket_suggestions(
                    header, row, document.source_path, table_index, index
                ),
                source_refs=(_anchor(document.source_path, locator),),
            )
        )
    if len(modules) < len(module_ids):
        modules.extend(
            _missing_modules(
                document.source_path,
                module_ids[len(modules) :],
                "syllabus row missing for declared module",
            )
        )
    return tuple(modules)


def _missing_modules(
    source_path: str, module_ids: list[str], reason: str
) -> tuple[ModuleMetadataProposal, ...]:
    return tuple(
        ModuleMetadataProposal(
            module_id=module_id,
            status="missing",
            slug_status="missing",
            slug_review_note=_anchored(reason, source_path, "syllabus structure"),
        )
        for module_id in module_ids
    )


def _modules(
    document: ExtractedSyllabusDocument, module_ids: list[str]
) -> tuple[ModuleMetadataProposal, ...]:
    if document.source_format == "docx":
        return _docx_modules(document, module_ids)
    return _mhtml_modules(document, module_ids)


def _format_unsupported(course_id: str, source_path: Path, reason: str) -> CourseMetadataProposal:
    return CourseMetadataProposal(
        course_id=course_id,
        source_path=source_path.as_posix(),
        extraction_status="format_unsupported",
        gaps=(
            GapEntry(
                kind="format_unsupported",
                severity="warning",
                path=source_path.as_posix(),
                message=reason,
            ),
        ),
    )


def build_module_metadata_proposal(
    course_root: Path,
    syllabus_path: Path,
    *,
    expected_learning_objective_count: int,
    required_title: str | None = None,
) -> CourseMetadataProposal:
    course = load_course(course_root)
    try:
        document = read_syllabus(syllabus_path)
    except SyllabusExtractionError as exc:
        return _format_unsupported(course.course_id, syllabus_path, str(exc))
    return build_module_metadata_proposal_from_document(
        course_root,
        document,
        expected_learning_objective_count=expected_learning_objective_count,
        required_title=required_title,
    )


def build_module_metadata_proposal_from_document(
    course_root: Path,
    document: ExtractedSyllabusDocument,
    *,
    expected_learning_objective_count: int,
    required_title: str | None = None,
) -> CourseMetadataProposal:
    course = load_course(course_root)
    module_ids = [module.module_id for module in course.modules]
    if not module_ids:
        modules_dir = course_root / "modules"
        module_ids = (
            sorted(path.name for path in modules_dir.iterdir() if path.is_dir())
            if modules_dir.exists()
            else []
        )
    objectives = _course_objectives(document)
    modules = _modules(document, module_ids)
    course_title = _course_title(document)
    sentinel_errors: list[str] = []
    if len(objectives) != expected_learning_objective_count:
        sentinel_errors.append(
            "course learning objective count did not match the expected sentinel"
        )
    if len(modules) != course.course.module_count_expected:
        sentinel_errors.append("module record count did not match course.yaml")
    if required_title is not None and (
        course_title is None or course_title.value != required_title
    ):
        sentinel_errors.append("course title did not match the required sentinel")
    if sentinel_errors:
        return _format_unsupported(
            course.course_id,
            Path(document.source_path),
            "; ".join(sentinel_errors),
        )
    if not objectives or not modules or all(module.status == "missing" for module in modules):
        return _format_unsupported(
            course.course_id,
            Path(document.source_path),
            "syllabus text was extracted but required metadata sentinels were missing",
        )
    return CourseMetadataProposal(
        course_id=course.course_id,
        source_path=document.source_path,
        extraction_status="verified",
        course_title=course_title,
        learner_profile=_learner_profile(document),
        course_learning_objectives=objectives,
        modules=modules,
    )


def render_module_metadata_yaml(proposal: CourseMetadataProposal) -> str:
    return yaml.safe_dump(proposal.model_dump(mode="json"), sort_keys=False, allow_unicode=True)


__all__ = [
    "AnchoredObjective",
    "AnchoredText",
    "CourseMetadataProposal",
    "ExtractedSyllabusDocument",
    "ModuleMetadataProposal",
    "SCHEMA_VERSION",
    "SourceAnchor",
    "SyllabusExtractionError",
    "build_module_metadata_proposal",
    "build_module_metadata_proposal_from_document",
    "load_extracted_document_yaml",
    "read_docx_syllabus",
    "read_mhtml_syllabus",
    "read_syllabus",
    "render_extracted_document_yaml",
    "render_module_metadata_yaml",
]
