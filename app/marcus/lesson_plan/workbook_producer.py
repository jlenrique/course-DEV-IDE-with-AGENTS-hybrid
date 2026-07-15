"""Workbook producer for the braid client deliverable (Braid Slice 1 / S2).

Audience: braid producers + reviewers consuming the 4th concrete
``ModalityProducer`` (after ``BlueprintProducer``). This module emits the
client **workbook companion** to the narrated deck — the dual-coding partner
that carries the transcript + a fuller narrative, exercises, and citations that
were deliberately kept OFF the glance-slides and out of heard-only narration.

Render path (DP3, binding):
    * **Markdown is canonical** — the citation-injection substrate S3 fills.
    * **DOCX is the client deliverable** — rendered FROM the same composed
      markdown model via ``python-docx`` (one source of truth; the deliverable
      cannot diverge from the canonical).
    * **PDF is DEFERRED to v-next** (named; requires a new on-disk dependency —
      pandoc/weasyprint/reportlab — none present). Never author "PDF/DOCX" as a
      single target.

Discipline notes:
    * Mirrors ``BlueprintProducer`` exactly for the deterministic skeleton:
      repo-root output containment, repo-relative asset paths, a default body
      renderer + an injectable seam, ``ProducedAsset`` return with
      ``fulfills=f"{plan_unit.unit_id}@{context.lesson_plan_revision}"``.
    * **Prose-delegation seam (Irene A4-3):** the Pass-2 transcript is the
      workbook's SCAFFOLD, not its text. The deterministic default embeds each
      segment verbatim under its ``segment_id`` AND emits an explicit
      ``REVOICE-REQUIRED`` marker so deixis-laden transcript text is never
      silently passed off as finished read-prose. An injectable
      ``prose_revoicer`` replaces the default body per segment without a
      producer-class change.
    * **L2 audit hook (FAIL-mode):** ``produce`` calls
      ``audit_numeric_provenance`` (read-only; the frozen-neck module is
      UNMODIFIED) and the producer asserts ``unsourced_numeric == 0`` at the
      workbook boundary (G1). The underlying module stays warn-only.
    * **id-uniqueness:** section_id / exercise_id binding boundary rejects
      duplicate ids (filed follow-on ``braid-collateral-id-uniqueness``).
    * **DP6 frozen-Gamma reuse:** slide-independent; the producer records a
      reuse stamp on the returned sidecar (workbook-only diff => reuse-OK).
"""

from __future__ import annotations

import re
import zipfile
from collections.abc import Callable, Iterable, Mapping, Sequence
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any, Final, Literal

import yaml
from docx import Document

from app.marcus.lesson_plan.collateral_spec import WorkbookSpec
from app.marcus.lesson_plan.modality_producer import ModalityProducer
from app.marcus.lesson_plan.prework_projection import PreWorkBrief, render_prework_markdown
from app.marcus.lesson_plan.produced_asset import ProducedAsset, ProductionContext
from app.marcus.lesson_plan.schema import PlanUnit
from app.specialists._shared.source_fidelity_audit import audit_numeric_provenance
from scripts.utilities.slide_production_gate import (
    fresh_gamma_required,
    reuse_stamp,
)

if TYPE_CHECKING:
    from app.marcus.lesson_plan.glossary_projection import GlossaryArticleBrief
    from app.marcus.lesson_plan.trends_projection import ResearchTrendsBrief

REPO_ROOT: Final[Path] = Path(__file__).resolve().parents[3]
DEFAULT_WORKBOOK_OUTPUT_ROOT: Final[str] = "_bmad-output/artifacts/workbooks"

# B8 (DOCX determinism): python-docx stamps wall-clock mtimes into the OPC zip
# members and writes ``dcterms:created/modified`` into ``docProps/core.xml``, so
# two renders of the SAME model produce different bytes (breaking reload-equality
# once the deliverable is captured trial-scoped). Pin both to a fixed epoch (the
# zip format's 1980-01-01 minimum) so the rendered DOCX is byte-deterministic.
_DETERMINISTIC_DOCX_ZIP_DATETIME: Final[tuple[int, int, int, int, int, int]] = (
    1980,
    1,
    1,
    0,
    0,
    0,
)
_DETERMINISTIC_DOCX_CORE_DT: Final[datetime] = datetime(1980, 1, 1, tzinfo=UTC)

# Prose-delegation review marker — the honesty guard that "transcript backbone"
# did NOT collapse into "workbook text" (mirrors BlueprintProducer's
# IRENE_REVIEW_MARKER discipline).
REVOICE_REQUIRED_MARKER: Final[str] = (
    "<!-- REVOICE-REQUIRED: writer (Paige/Sophia) -> self-contained read-prose; Irene validates -->"
)
HUMAN_REVIEW_SECTION_HEADING: Final[str] = "## Human Review Checkpoint"
IRENE_REVIEW_MARKER: Final[str] = "- [ ] Irene workbook review complete"

# Stable per-segment anchor format (AC-5 line-level traceability).
SEGMENT_ANCHOR_PREFIX: Final[str] = "segment:"

# G4 close-out note — this story makes NO reading-path generalization claim.
NO_READING_PATH_HALO_NOTE: Final[str] = (
    "<!-- G4: no reading-path halo; this workbook advances no "
    "fresh-naive-holdout / reading-path generalization claim -->"
)


# ``prose_revoicer`` seam: given a segment's id + verbatim narration text,
# returns the body prose to emit for that segment. The default keeps the
# verbatim text and appends a REVOICE-REQUIRED marker.
ProseRevoicer = Callable[[str, str], str]


class WorkbookScopeError(ValueError):
    """Raised when a plan unit / spec is not eligible for workbook production."""


class WorkbookFidelityError(ValueError):
    """Raised when an honesty-gate assertion fails at the workbook boundary."""


class DuplicateCollateralIdError(ValueError):
    """Raised when a section_id / exercise_id duplicate is detected.

    The asset->objective binding boundary enforces id-uniqueness (filed
    follow-on ``braid-collateral-id-uniqueness``).
    """


# ---------------------------------------------------------------------------
# Transcript-segment backbone (the run's frozen deck segment-manifest input)
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class TranscriptSegment:
    """One Pass-2 transcript segment read from the frozen segment manifest.

    Read directly from the raw manifest YAML (the on-disk shape is richer than
    the strict authoring projection); only the fields the workbook needs are
    projected here.
    """

    segment_id: str
    narration_text: str
    visual_file: str | None
    visual_references: tuple[dict[str, Any], ...]
    source_ref: str
    # Absolute, on-disk path to the figure image, resolved against the manifest's
    # bundle directory at load time (the manifest ``visual_file`` is
    # bundle-relative, e.g. ``gamma-export/...png`` — it does NOT resolve against
    # the repo root). ``None`` when the segment has no figure.
    visual_file_abs: str | None = None


def load_transcript_segments(manifest_path: str | Path) -> tuple[TranscriptSegment, ...]:
    """Load Pass-2 transcript segments from a segment-manifest YAML.

    Reads the raw manifest (offline, read-only). Each segment's ``id`` is the
    transcript-scaffold anchor (AC-5); ``narration_text`` is the VO/transcript
    text; ``visual_file`` + ``visual_references`` feed the figure embed (AC-7).
    ``source_ref`` defaults to the segment id when the manifest does not carry
    an explicit one (the thin S2 source set; S3 fills real refs).
    """
    path = Path(manifest_path)
    bundle_dir = path.resolve().parent
    data = yaml.safe_load(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict) or "segments" not in data:
        raise WorkbookScopeError(f"segment manifest {path} missing top-level 'segments' list")
    raw_segments = data["segments"]
    if not isinstance(raw_segments, list) or not raw_segments:
        raise WorkbookScopeError(f"segment manifest {path} has no segments")

    out: list[TranscriptSegment] = []
    seen_ids: set[str] = set()
    for raw in raw_segments:
        seg_id = raw.get("id")
        if not isinstance(seg_id, str) or not seg_id:
            raise WorkbookScopeError(f"segment manifest {path} has a segment with no 'id'")
        if seg_id in seen_ids:
            raise DuplicateCollateralIdError(
                f"duplicate transcript segment id {seg_id!r} in {path}"
            )
        seen_ids.add(seg_id)
        vrefs = raw.get("visual_references") or []
        visual_file = raw.get("visual_file")
        # Resolve the figure path against the BUNDLE directory (the manifest's
        # parent), NOT the repo root — the manifest path is bundle-relative.
        visual_file_abs: str | None = None
        if isinstance(visual_file, str) and visual_file:
            candidate = Path(visual_file)
            resolved = candidate if candidate.is_absolute() else bundle_dir / candidate
            if resolved.exists():
                visual_file_abs = str(resolved)
        out.append(
            TranscriptSegment(
                segment_id=seg_id,
                narration_text=str(raw.get("narration_text", "")),
                visual_file=visual_file,
                visual_references=tuple(v for v in vrefs if isinstance(v, dict)),
                source_ref=str(raw.get("source_ref") or seg_id),
                visual_file_abs=visual_file_abs,
            )
        )
    return tuple(out)


# ---------------------------------------------------------------------------
# S1 / S6 produce-time content briefs (production inputs, NOT spec schema)
# ---------------------------------------------------------------------------
#
# These are the SOURCE content the producer composes into the deliverable —
# siblings of ``segments`` / ``source_text`` / ``citations`` (all already
# produce-time inputs, not ``WorkbookSpec`` fields). The spec carries the
# STRUCTURE + bindings + intent (section -> learning_objective_id, exercise ->
# answer_key_source_ref); the run supplies the lesson CONTENT (the objective
# statements, the worked answers, the bibliography) that the producer renders.
# Keeping them here avoids a CollateralSpec schema bump for a producer-only
# composition story (design §7 gaps 1/4/5 = composition, not new substrate).


@dataclass(frozen=True)
class LearningObjectiveBrief:
    """One learning objective surfaced in the workbook's S1 section.

    ``objective_id`` MUST match a section's ``learning_objective_id`` (the
    asset->objective binding). The producer asserts no-orphan / no-phantom
    against the section bindings when objectives are supplied.
    """

    objective_id: str
    bloom_level: str
    statement: str


@dataclass(frozen=True)
class FurtherReadingEntry:
    """A corpus-native / required-reading citation rendered into S6.

    Deck-traceable references (CMS, JAMA, the required-reading article, the
    intro video). ``source_ref`` is the traceable reference the G2 citation
    audit resolves against this run's manifest; ``locator`` is the human-facing
    URL / publication string.
    """

    citation_id: str
    title: str
    source_ref: str
    locator: str | None = None
    supports_segment_id: str | None = None


@dataclass(frozen=True)
class ResearchEntry:
    """A live-research cited entry (braid S3 Texas/Scite leg) rendered into S6.

    Each carries a real ``source_id`` DOI; the producer renders the link as
    ``https://doi.org/{source_id}``. Empty in a v1 artifact whose research leg
    is not yet live (rendered as an explicit "deferred" note, never fabricated).

    R4 credibility fields are additive (tier / peer-review / provenance /
    triangulation) — projected when present; absent on older envelopes.
    """

    citation_id: str
    title: str
    source_ref: str
    provider: str
    source_id: str
    source_hash: str | None = None
    supports_segment_id: str | None = None
    evidence_hierarchy_tier: str | None = None
    peer_reviewed: bool | None = None
    provider_provenance: tuple[str, ...] | None = None
    triangulation_status: str | None = None
    reliability_score: float | None = None


# ---------------------------------------------------------------------------
# id-uniqueness at the asset->objective binding boundary
# ---------------------------------------------------------------------------


def assert_unique_collateral_ids(spec: WorkbookSpec) -> None:
    """Reject duplicate section_id / exercise_id / goal binding ids.

    This is the agreed enforcement point for the filed follow-on
    ``braid-collateral-id-uniqueness``: the producer binds asset -> objective by
    these ids, so a duplicate would silently collide two bindings.
    """
    section_ids: list[str] = []
    exercise_ids: list[str] = []
    for section in spec.sections:
        section_ids.append(section.section_id)
        for exercise in section.exercises:
            exercise_ids.append(exercise.exercise_id)
    _reject_dupes(section_ids, label="section_id")
    _reject_dupes(exercise_ids, label="exercise_id")


def _reject_dupes(ids: Sequence[str], *, label: str) -> None:
    seen: set[str] = set()
    for value in ids:
        if value in seen:
            raise DuplicateCollateralIdError(
                f"duplicate {label} {value!r} at the workbook binding boundary"
            )
        seen.add(value)


# ---------------------------------------------------------------------------
# L2 audit FAIL-mode gate wrapper (the frozen-neck module stays UNMODIFIED)
# ---------------------------------------------------------------------------


def audit_workbook_numeric_fidelity(
    workbook_body_text: str,
    source_text: str,
    *,
    research_supplements: set[str] | None = None,
) -> dict:
    """Run the L2 numeric audit in FAIL-mode at the workbook boundary (G1).

    Calls the frozen-neck ``audit_numeric_provenance`` READ-ONLY and converts
    the warn-only return into a FAIL-mode gate: ``unsourced_numeric == 0`` must
    hold or :class:`WorkbookFidelityError` is raised. An un-auditable
    zero-denominator FAIL is ALSO raised — but only when the workbook body itself
    carries numeric tokens (a numeral that cannot be cleared against a
    numeral-free source); a genuinely numeral-free workbook is a clean pass. The
    underlying module is unmodified — the FAIL decision lives here.

    HONESTY BOUNDARY — word-form numerals are NOT gated (named gap, do NOT
    over-claim): the L2 figure tokenizer is SYMBOL-ONLY (``42%``, ``$7`` — the
    frozen-neck ``_FIGURE_RE``). Word-form numerals ("forty-two percent",
    "42 percent") are invisible to this gate, so a word-form numeral assertion
    passes un-audited. This is a deliberate, named limitation, not a guarantee of
    full numeric coverage. Filed as deferred follow-on
    ``braid-workbook-wordform-numeral-gap``; not solved here.

    Returns the report (attached to the run record / sidecar).
    """
    report = audit_numeric_provenance(
        workbook_body_text,
        source_text,
        research_supplements=research_supplements,
    )
    # The frozen-neck returns status="FAIL" (count==0) when EITHER side has zero
    # numeric tokens — the audit is UN-AUDITABLE (zero-denominator guard F2), not
    # a clean pass. The credibility risk this gate exists to catch is: the
    # workbook body ASSERTS a numeral but the source set carries NONE, so the
    # numeral cannot be cleared. That case (``narration_token_count > 0`` under a
    # FAIL) must RAISE — the OLD wrapper let it slip as "0 unsourced". When the
    # workbook body itself carries no numeric tokens, there is nothing to gate and
    # the (genuinely numeral-free) workbook passes; a FAIL there is a non-event,
    # not a violation, so we do NOT over-claim a fidelity breach.
    if report.get("status") == "FAIL" and report.get("narration_token_count", 0) > 0:
        raise WorkbookFidelityError(
            "G1 FAIL-mode: the L2 numeric audit is un-auditable for a workbook that "
            "ASSERTS numerals (zero-denominator: "
            f"{report.get('reason', 'source carries no numeric tokens')}). "
            "Treating un-auditable as a gate failure at the workbook boundary — a "
            "workbook numeral cannot be cleared against a numeral-free source set."
        )
    unsourced = report["buckets"]["unsourced_numeric"]["count"]
    if unsourced != 0:
        raise WorkbookFidelityError(
            f"G1 FAIL-mode: workbook body carries {unsourced} unsourced numeric "
            "assertion(s); every numeral must trace to the source set or a "
            "declared research supplement"
        )
    return report


# ---------------------------------------------------------------------------
# Citation-fidelity wiring (G2) — thin in S2; the assertion path is the deliverable
# ---------------------------------------------------------------------------


def audit_citation_fidelity(
    citations: Iterable[dict[str, str]],
    source_ref_manifest: dict[str, str],
) -> dict:
    """Assert every workbook citation resolves to a real ``source_ref`` (G2).

    Each citation is a ``{"source_ref": ...}`` mapping; the manifest maps
    ``source_ref -> source-hash``. FAIL-mode: ``unsourced_citations == 0``.
    Returns a report with the citation->source_ref->hash manifest attached.
    """
    citations = list(citations)
    unresolved: list[str] = []
    resolved: list[dict[str, str]] = []
    for citation in citations:
        ref = citation.get("source_ref", "")
        source_hash = source_ref_manifest.get(ref)
        if not ref or source_hash is None:
            unresolved.append(ref)
        else:
            resolved.append({"source_ref": ref, "source_hash": source_hash})
    report = {
        "buckets": {
            "unsourced_citations": {"count": len(unresolved), "refs": unresolved},
        },
        "resolved": resolved,
        "manifest": dict(source_ref_manifest),
    }
    if unresolved:
        raise WorkbookFidelityError(
            f"G2 FAIL-mode: {len(unresolved)} workbook citation(s) do not resolve "
            f"to a real source_ref: {unresolved!r}"
        )
    return report


_DEEP_DIVE_CITE_MARKER_RE: Final[re.Pattern[str]] = re.compile(
    r"\[(ask-a-cite-[0-9]{3})\]"
)


def deep_dive_g2_citation_entries(
    markdown: str, deep_dive_review: Any
) -> list[dict[str, str]]:
    """G2 deep-dive leg (37.2b, non-tautological by construction).

    PROSE side: every inline ``[ask-a-cite-###]`` marker parsed from the
    RENDERED workbook markdown. STRUCTURED side: the citation_id -> source_ref
    map built from the contribution's USED pool rows. A marker whose id has no
    used row (prose/structure divergence, a stray marker on a degraded or
    absent contribution, or a corrupted render) maps to a deliberately
    unresolvable ref and therefore FAILS ``audit_citation_fidelity``. What this
    leg does NOT prove: structured used-rows lacking prose markers (row g of
    the A2 gate owns that direction) and semantic claim<->source support
    (operator WARN).
    """
    rows_by_citation: dict[str, str] = {}
    if deep_dive_review is not None:
        result = getattr(deep_dive_review, "deep_dive_enrichment", None)
        if result is not None and result.status == "enriched":
            from app.marcus.lesson_plan.deep_dive_enrichment import (  # noqa: PLC0415
                used_pool_rows,
            )

            rows_by_citation = {
                row.citation_id: row.source_ref for row in used_pool_rows(result)
            }
    return [
        {
            "source_ref": rows_by_citation.get(
                marker_id, f"deep-dive-citation-unresolved:{marker_id}"
            )
        }
        for marker_id in sorted(set(_DEEP_DIVE_CITE_MARKER_RE.findall(markdown)))
    ]


def assert_exercise_fidelity(spec: WorkbookSpec) -> None:
    """G3: every exercise carries a Bloom level + a present answer-key source_ref."""
    for section in spec.sections:
        for exercise in section.exercises:
            if not exercise.bloom_level:
                raise WorkbookFidelityError(
                    f"G3: exercise {exercise.exercise_id!r} missing a Bloom level"
                )
            if not (exercise.answer_key_source_ref or "").strip():
                raise WorkbookFidelityError(
                    f"G3: exercise {exercise.exercise_id!r} answer key has no resolvable source_ref"
                )


def assert_learning_objective_bindings(
    spec: WorkbookSpec,
    learning_objectives: Sequence[LearningObjectiveBrief],
) -> None:
    """S1 fidelity: no orphan objective, no phantom objective.

    When the run surfaces ``learning_objectives`` (the S1 render input), every
    listed objective MUST be bound by >=1 section's ``learning_objective_id``
    (no ORPHAN objective listed that nothing serves) AND every section's
    ``learning_objective_id`` MUST appear in the listed objectives (no PHANTOM
    binding pointing at an objective the workbook never states). When no
    objectives are supplied the check is a no-op (non-regression: the producer
    then renders a thin id-only objectives list derived from the bindings).
    """
    if not learning_objectives:
        return
    bound_ids = {section.learning_objective_id for section in spec.sections}
    listed_ids = {obj.objective_id for obj in learning_objectives}
    orphans = sorted(listed_ids - bound_ids)
    if orphans:
        raise WorkbookFidelityError(
            f"S1: orphan learning objective(s) listed but bound by no section: {orphans!r}"
        )
    phantoms = sorted(bound_ids - listed_ids)
    if phantoms:
        raise WorkbookFidelityError(
            "S1: section(s) bind a learning_objective_id with no stated "
            f"objective (phantom binding): {phantoms!r}"
        )


# ---------------------------------------------------------------------------
# Default prose revoicer (deterministic; honesty marker, no silent passthrough)
# ---------------------------------------------------------------------------


def default_prose_revoicer(segment_id: str, narration_text: str) -> str:
    """Embed the verbatim transcript text + the REVOICE-REQUIRED honesty marker.

    The default NEVER asserts deixis-laden transcript text is finished
    read-prose: the verbatim segment is preserved (the scaffold) and the marker
    flags that a writer (Paige/Sophia) must re-voice it into self-contained
    read-prose, with Irene validating behavioral intent.
    """
    return "\n".join(
        [
            REVOICE_REQUIRED_MARKER,
            f"> {narration_text.strip()}",
        ]
    )


# ---------------------------------------------------------------------------
# Markdown composition (canonical) — single source of truth for DOCX
# ---------------------------------------------------------------------------


@dataclass
class _ComposedDoc:
    """The composed canonical model: ordered (heading_level, heading, body) blocks.

    DOCX renders from THIS model so the deliverable cannot diverge from the
    canonical markdown.
    """

    title: str
    blocks: list[tuple[int, str, str]] = field(default_factory=list)
    # (abs_image_path_or_None, caption, source_ref) — abs path is the on-disk
    # image resolved against the bundle dir (None when the figure file is absent,
    # so the DOCX renderer skips the picture but still records the caption).
    figures: list[tuple[str | None, str, str]] = field(default_factory=list)


def _segment_anchor(segment_id: str) -> str:
    return f"<!-- {SEGMENT_ANCHOR_PREFIX}{segment_id} -->"


def compose_workbook(
    plan_unit: PlanUnit,
    context: ProductionContext,
    spec: WorkbookSpec,
    segments: Sequence[TranscriptSegment],
    *,
    prose_revoicer: ProseRevoicer,
    workbook_title: str | None = None,
    learning_objectives: Sequence[LearningObjectiveBrief] = (),
    answer_keys: Mapping[str, str] | None = None,
    further_reading: Sequence[FurtherReadingEntry] = (),
    research_entries: Sequence[ResearchEntry] = (),
    research_empty_reason: str | None = None,
    research_omitted_note: str | None = None,
    glossary_articles: Sequence[GlossaryArticleBrief] = (),
    glossary_empty_reason: str | None = None,
    glossary: Any | None = None,
    research_trends: ResearchTrendsBrief | None = None,
    pre_work: PreWorkBrief | None = None,
    encounter_mode: Literal["recorded", "live"] = "recorded",
    render_profile: Literal["legacy", "presentation_support"] = "legacy",
    lo_overlay_loss_note: str | None = None,
    deep_dive_review: Any | None = None,
) -> _ComposedDoc:
    """Compose the canonical workbook model from the spec + transcript backbone.

    Sections composed (design §2 v1 cut S0-S7 + W2 glossary + W3 trends):
      S0 Overview, S1 Learning Objectives, S2 Transcript-narrative (re-voiced
      deferred depth) + Depth-delta narrative, S3 Transcript of Record (literal
      verbatim), S4 Figures, S5 Exercises + Answer Key, W2 Research Glossary
      (encyclopedia articles from wrangled research), S6 References /
      further-reading, W3 Research Trends + hot-topics, S7 Human Review footer
      (added by the renderers).
    """
    # Lazy import: glossary_projection → research_packet → workbook_enrichment
    # → this module (avoid circular import at package load).
    from app.marcus.lesson_plan.glossary_projection import (  # noqa: PLC0415
        render_glossary_markdown,
    )
    from app.marcus.lesson_plan.trends_projection import (  # noqa: PLC0415
        ResearchTrendsBrief as _TrendsBrief,
    )
    from app.marcus.lesson_plan.trends_projection import (
        render_trends_markdown,
    )

    answer_keys = dict(answer_keys or {})
    # The DISPLAY title is a human-readable label distinct from the open-id
    # ``event_type`` slug (which reads as a hyphenated 80-char id on the H1).
    # Fall back to ``event_type`` when no display title is supplied.
    title_label = (workbook_title or "").strip() or plan_unit.event_type
    doc = _ComposedDoc(title=f"Workbook: {title_label}")

    if pre_work is not None:
        label = {
            "recorded": "Before you watch the recorded presentation",
            "live": "Before the live presentation",
        }[encounter_mode]
        doc.blocks.append((2, "Pre-work", label))
        rendered = render_prework_markdown(pre_work).rstrip().splitlines()
        current_heading: str | None = None
        body: list[str] = []
        for line in rendered:
            if line.startswith("## "):
                if current_heading is not None:
                    doc.blocks.append((2, current_heading, "\n".join(body).strip()))
                current_heading, body = line[3:], []
            else:
                body.append(line)
        if current_heading is not None:
            doc.blocks.append((2, current_heading, "\n".join(body).strip()))

    # --- S0 Overview ---
    doc.blocks.append(
        (
            2,
            "Overview",
            "\n".join(
                [
                    f"Unit ID: `{plan_unit.unit_id}`",
                    f"Fulfills target: `{plan_unit.unit_id}@{context.lesson_plan_revision}`",
                    "",
                    (
                        "This presentation-support workbook carries pre-work, "
                        "guided practice, exercises, and citations."
                        if render_profile == "presentation_support"
                        else "This workbook is the read-in-depth companion to the narrated "
                        "deck: it carries the transcript, the fuller narrative deferred "
                        "off the glance-slides, the exercises, and the citations."
                    ),
                    "",
                    (
                        "Use this workbook before and alongside the presentation: "
                        "complete the pre-work, read the Deep Dive, work the "
                        "self-check exercises, and follow the cited sources. The "
                        "Deep Dive section carries the traceable read-prose for "
                        "this lesson; when the run's cited research pool authored "
                        "enrichment, its sentences cite that pool inline."
                        if render_profile == "presentation_support"
                        else "How to use this workbook with the deck: watch the "
                        "glance-deck for the perception-tuned claim per card, then "
                        "read here for the deferred depth, work the self-check "
                        "exercises, and follow the further reading into primary "
                        "sources. The deck carries the *glance* channel; this "
                        "workbook carries the *read* channel (dual-coding partners)."
                    ),
                ]
            ),
        )
    )

    # --- S1 Learning Objectives (render the bindings; never surfaced before) ---
    objective_lines: list[str] = []
    if learning_objectives:
        for obj in learning_objectives:
            serving = sorted(
                section.section_id
                for section in spec.sections
                if section.learning_objective_id == obj.objective_id
            )
            serving_note = (
                f" (served by section(s): {', '.join(f'`{s}`' for s in serving)})"
                if serving
                else ""
            )
            objective_lines.append(
                f"- **`{obj.objective_id}`** — _{obj.bloom_level}_: {obj.statement}{serving_note}"
            )
    else:
        # Non-regression thin render: derive the distinct objective ids bound by
        # sections (the bindings exist even when no statements are supplied).
        seen: set[str] = set()
        for section in spec.sections:
            oid = section.learning_objective_id
            if oid in seen:
                continue
            seen.add(oid)
            objective_lines.append(f"- `{oid}` (served by section `{section.section_id}`)")
    # Q1: a visible enrichment-overlay loss note (degrade-with-record; never a
    # silent placeholder swap). Numeral-free + outside the AC-8 narrative blocks.
    if lo_overlay_loss_note:
        objective_lines.append("")
        objective_lines.append(f"> _Enrichment overlay loss: {lo_overlay_loss_note}_")
    doc.blocks.append((2, "Learning Objectives", "\n".join(objective_lines).rstrip()))

    # --- Deep Dive (37.2b): cited, pool-grounded read-prose at the enrichment
    # seam. Deterministic-consume and model-free: the section renders the
    # persisted 07W.3 contribution (enriched prose with inline
    # ``[ask-a-cite-###]`` markers) or the honest typed-loss note — never a
    # silent absence. Presentation-support profile only.
    if render_profile == "presentation_support":
        from app.marcus.lesson_plan.deep_dive_enrichment import (  # noqa: PLC0415
            render_deep_dive_markdown,
        )

        doc.blocks.append((2, "Deep Dive", render_deep_dive_markdown(deep_dive_review)))

    # --- S2 Transcript-narrative (the prose-delegation seam + segment backbone) ---
    if render_profile == "legacy":
        transcript_lines: list[str] = []
        for seg in segments:
            transcript_lines.append(_segment_anchor(seg.segment_id))
            transcript_lines.append(f"#### {seg.segment_id}")
            transcript_lines.append(prose_revoicer(seg.segment_id, seg.narration_text))
            transcript_lines.append("")
        doc.blocks.append((2, "Transcript-narrative", "\n".join(transcript_lines).rstrip()))

    # --- Section depth-deltas (the dual-coding legitimization of the tight VO) ---
    section_lines: list[str] = []
    for section in spec.sections:
        section_lines.append(f"### {section.title}")
        section_lines.append(
            f"Objective: `{section.learning_objective_id}` (section `{section.section_id}`)"
        )
        section_lines.append(f"Depth deferred off the slide: {section.depth_delta.deferred_depth}")
        if section.narrative_intent:
            section_lines.append(section.narrative_intent)
        section_lines.append("")
    doc.blocks.append((2, "Depth-delta narrative", "\n".join(section_lines).rstrip()))

    # --- S3 Transcript of Record (literal verbatim; distinct from the re-voiced
    # S2 explainer above). The clean record of what was heard, segment by
    # segment — NO REVOICE-REQUIRED marker, NO depth framing. This separates the
    # verbatim transcript (record) from the re-voiced read-prose (S2), the
    # design's S2/S3 split (§7 gap 2). Excluded from the AC-8 superset (it IS
    # the VO of record, not depth beyond it).
    if render_profile == "legacy":
        record_lines: list[str] = []
        for seg in segments:
            record_lines.append(f"#### {seg.segment_id} (verbatim)")
            record_lines.append(f"> {seg.narration_text.strip()}")
            record_lines.append("")
        doc.blocks.append((2, "Transcript of Record", "\n".join(record_lines).rstrip()))

    # --- Figures (image + caption + source_ref) ---
    figure_lines: list[str] = []
    for seg in segments:
        if not seg.visual_file or not seg.visual_references:
            continue
        vref = seg.visual_references[0]
        caption = str(vref.get("description") or vref.get("element") or seg.segment_id)
        source_ref = seg.source_ref
        # MD carries the readable bundle-relative ``![caption](path)`` reference;
        # the DOCX renderer embeds the BUNDLE-resolved absolute path (or skips if
        # the file is absent).
        figure_lines.append(_segment_anchor(seg.segment_id))
        figure_lines.append(f"![{caption}]({seg.visual_file})")
        figure_lines.append(f"*Figure — {caption}* (source_ref: `{source_ref}`)")
        figure_lines.append("")
        doc.figures.append((seg.visual_file_abs, caption, source_ref))
    if render_profile == "legacy":
        doc.blocks.append((2, "Figures", "\n".join(figure_lines).rstrip()))
    else:
        doc.figures.clear()

    # --- Exercises (Bloom level + source-grounded answer key) ---
    exercise_lines: list[str] = []
    for section in spec.sections:
        for exercise in section.exercises:
            exercise_lines.append(f"### Exercise `{exercise.exercise_id}`")
            exercise_lines.append(f"Bloom level: **{exercise.bloom_level}**")
            exercise_lines.append(exercise.prompt_intent)
            exercise_lines.append(f"Answer key (source_ref: `{exercise.answer_key_source_ref}`)")
            exercise_lines.append("")
    doc.blocks.append((2, "Exercises", "\n".join(exercise_lines).rstrip()))

    # --- S5b Answer Key (worked answers, source-grounded; G1/G3). The worked
    # answer prose is a produce-time input (the corpus "Correct Answer" line);
    # the spec carries only the answer_key_source_ref slot. When no worked answer
    # is supplied the producer emits an explicit pending note + the source_ref
    # (never a fabricated answer). ---
    answer_lines: list[str] = []
    for section in spec.sections:
        for exercise in section.exercises:
            answer_lines.append(f"### Answer — `{exercise.exercise_id}`")
            worked = answer_keys.get(exercise.exercise_id, "").strip()
            if worked:
                answer_lines.append(worked)
            else:
                answer_lines.append(
                    "*(worked answer pending writer; grounded by the "
                    f"source_ref below — source_ref: `{exercise.answer_key_source_ref}`)*"
                )
            answer_lines.append(f"Answer key source_ref: `{exercise.answer_key_source_ref}`")
            answer_lines.append("")
    doc.blocks.append((2, "Answer Key", "\n".join(answer_lines).rstrip()))

    # --- W2 Research Glossary. 39.1 (the learner-deliverable path): the
    # TERM-KEYED projection renders one entry per bolded Deep Dive term with
    # headword-identity association (AC-A5/A-6) — supplied as ``glossary``.
    # The _act intake supplies a projection on EVERY profile (legacy included;
    # its 07W.3 disk read is profile-agnostic — P7 legacy honesty), so a
    # legacy-profile run with an activated contribution renders real entries,
    # never a false absent-reason. Only callers that pass ``glossary_articles``
    # WITHOUT a projection (the frozen W2/W4 evidence scripts) take the
    # pre-39.1 article render below. Placed after Answer Key and before
    # References so W3 trends can stack after References later.
    if glossary is not None:
        from app.marcus.lesson_plan.glossary_projection import (  # noqa: PLC0415
            render_glossary_projection_markdown,
        )

        glossary_body = render_glossary_projection_markdown(glossary)
    else:
        glossary_body = render_glossary_markdown(
            glossary_articles,
            empty_reason=glossary_empty_reason,
        )
    doc.blocks.append((2, "Research Glossary", glossary_body))

    # --- S6 References / Further reading (rendered bibliography). Three tiers:
    # (a) per-segment source_ref traceability lines (the thin substrate),
    # (b) corpus-native + required-reading citations, (c) live-research DOI'd
    # entries (https://doi.org/{source_id}); an empty research set renders an
    # explicit deferred note, never a fabricated DOI (design §4). ---
    reference_lines = []
    reference_lines.append("#### Segment source-ref trace")
    for seg in segments:
        reference_lines.append(f"- `{seg.segment_id}` -> source_ref `{seg.source_ref}`")
    reference_lines.append("")
    reference_lines.append("#### Cited references & further reading")
    if further_reading:
        for entry in further_reading:
            supports = (
                f" — supports `{entry.supports_segment_id}`" if entry.supports_segment_id else ""
            )
            if entry.locator:
                reference_lines.append(
                    f"- [{entry.title}]({entry.locator}) "
                    f"(source_ref: `{entry.source_ref}`){supports}"
                )
            else:
                reference_lines.append(
                    f"- {entry.title} (source_ref: `{entry.source_ref}`){supports}"
                )
    else:
        reference_lines.append("- *(no corpus-native references supplied for this artifact)*")
    reference_lines.append("")
    reference_lines.append("#### Live-research cited entries (DOI)")
    if research_entries:
        for entry in research_entries:
            # F-2604: render supports_segment_id ONLY when present on the entry
            # (never inferred — that is the deferred semantic-audit arc).
            supports = (
                f" — supports `{entry.supports_segment_id}`" if entry.supports_segment_id else ""
            )
            credibility = ""
            if entry.evidence_hierarchy_tier:
                peer = (
                    "peer-reviewed"
                    if entry.peer_reviewed
                    else "not peer-reviewed"
                    if entry.peer_reviewed is False
                    else "peer-review unknown"
                )
                prov = (
                    ",".join(entry.provider_provenance)
                    if entry.provider_provenance
                    else entry.provider
                )
                tri = entry.triangulation_status or "none"
                score = (
                    f", reliability={entry.reliability_score:.2f}"
                    if entry.reliability_score is not None
                    else ""
                )
                credibility = (
                    f", tier={entry.evidence_hierarchy_tier}, {peer}, "
                    f"provenance={prov}, triangulation={tri}{score}"
                )
            reference_lines.append(
                f"- {entry.title}. https://doi.org/{entry.source_id} "
                f"(provider: {entry.provider}, source_ref: `{entry.source_ref}`, "
                f"citation_id: `{entry.citation_id}`{credibility}){supports}"
            )
    else:
        # D4 / F-2601: zero rows => recorded explicitly-empty WITH reason (mirror
        # S6 D4). NEVER a fabricated DOI, and NOT the retired "deferred" note.
        reason = research_empty_reason or (
            "no cited research entries recorded for this run; recorded "
            "explicitly-empty (corpus-native references above carry the citations)"
        )
        reference_lines.append(f"- *({reason})*")
    # Remediation item-3 (DOI honesty): entries whose DOI was malformed/absent are
    # excluded from the rendered list above; record their omission with visible
    # provenance (never a bare/broken ``https://doi.org/`` link).
    if research_omitted_note:
        reference_lines.append(f"- *({research_omitted_note})*")
    # --- 37.2b resolvability floor (A8): every inline ``ask-a-cite-###`` marker
    # in the Deep Dive prose resolves to a rendered reference entry carrying its
    # citation_id, DOI, tier, and provenance. Full References assemble/dedupe/
    # render ownership stays with Story 37.5.
    if render_profile == "presentation_support":
        from app.marcus.lesson_plan.deep_dive_enrichment import (  # noqa: PLC0415
            render_deep_dive_reference_lines,
        )

        deep_dive_reference_lines = list(
            render_deep_dive_reference_lines(deep_dive_review)
        )
        # 39.1 AC-A9 (W2/M1): APPEND glossary-covered rows into the SAME
        # cited-entries block, deduped by citation_id against the
        # enrichment-emitted lines (a citation cited by both the deep dive and
        # a glossary entry appears ONCE). Resolvability floor only — the 37.5
        # References assemble/dedupe ownership fence stands.
        glossary_reference_line_items: list[str] = []
        if glossary is not None:
            from app.marcus.lesson_plan.glossary_projection import (  # noqa: PLC0415
                glossary_reference_lines,
            )

            emitted_ids: tuple[str, ...] = ()
            review_result = getattr(deep_dive_review, "deep_dive_enrichment", None)
            if review_result is not None and review_result.status == "enriched":
                emitted_ids = tuple(review_result.gate.used_citation_ids)
            glossary_reference_line_items = list(
                glossary_reference_lines(glossary, exclude_citation_ids=emitted_ids)
            )
        if deep_dive_reference_lines or glossary_reference_line_items:
            reference_lines.append("")
            reference_lines.append("#### Deep Dive cited entries (Ask-A, DOI)")
            reference_lines.extend(deep_dive_reference_lines)
            reference_lines.extend(glossary_reference_line_items)
    doc.blocks.append((2, "References", "\n".join(reference_lines).rstrip()))

    # --- W3 Research Trends + hot-topics (after References backmatter).
    trends_brief = research_trends
    if trends_brief is None:
        trends_brief = _TrendsBrief(
            trends=(),
            hot_topics=(),
            known_losses=(),
            empty_reason=(
                "no research-trends brief supplied for this artifact; recorded explicitly-empty"
            ),
        )
    doc.blocks.append((2, "Research Trends", render_trends_markdown(trends_brief)))

    return doc


def render_markdown(doc: _ComposedDoc) -> str:
    """Render the canonical Markdown from the composed model."""
    lines: list[str] = [f"# {doc.title}", ""]
    for level, heading, body in doc.blocks:
        lines.append(f"{'#' * level} {heading}")
        lines.append("")
        if body:
            lines.append(body)
        lines.append("")
    lines.append(HUMAN_REVIEW_SECTION_HEADING)
    lines.append(IRENE_REVIEW_MARKER)
    lines.append(NO_READING_PATH_HALO_NOTE)
    lines.append("")
    return "\n".join(lines)


def render_docx_from_model(doc: _ComposedDoc, output_path: Path) -> None:
    """Render the DOCX deliverable FROM the same composed model (not parallel).

    The DOCX heading text is derived from ``doc.blocks`` (the same source the
    canonical Markdown renders from) so the deliverable cannot diverge.
    """
    document = Document()
    document.add_heading(doc.title, level=0)
    for level, heading, body in doc.blocks:
        document.add_heading(heading, level=min(level, 4))
        if body:
            _render_docx_body(document, body)
    # Embed figures (AC-7): real image part (word/media/*) + caption + source_ref.
    for fig_path, caption, source_ref in doc.figures:
        if fig_path and Path(fig_path).exists():
            document.add_picture(str(fig_path))
        document.add_paragraph(f"Figure — {caption} (source_ref: {source_ref})")
    document.add_heading("Human Review Checkpoint", level=2)
    _render_docx_body(document, IRENE_REVIEW_MARKER)
    # B8: pin the wall-clock core-property timestamps BEFORE save (removes the
    # volatile ``dcterms:created/modified`` from ``docProps/core.xml``).
    document.core_properties.created = _DETERMINISTIC_DOCX_CORE_DT
    document.core_properties.modified = _DETERMINISTIC_DOCX_CORE_DT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    # B8: pin the OPC zip member mtimes AFTER save (python-docx stamps wall-clock
    # into each member via ``zipfile.writestr``) so the bytes are deterministic.
    _pin_docx_zip_determinism(output_path)


def _pin_docx_zip_determinism(path: Path) -> None:
    """Rewrite the DOCX zip with a fixed member timestamp for byte-determinism.

    python-docx writes each OPC part with ``zipfile.writestr``, which stamps the
    current wall-clock into the member ``date_time``; two renders of the same
    model therefore differ byte-for-byte. Rewrite every member with the pinned
    epoch (preserving member order, names, external attrs, and DEFLATE
    compression) so reload-equality holds once the deliverable is captured.
    """
    with zipfile.ZipFile(path, "r") as zin:
        members = [
            (info.filename, info.external_attr, zin.read(info.filename))
            for info in zin.infolist()
        ]
    tmp_path = path.with_name(path.name + ".det.tmp")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for filename, external_attr, data in members:
            member = zipfile.ZipInfo(filename, date_time=_DETERMINISTIC_DOCX_ZIP_DATETIME)
            member.compress_type = zipfile.ZIP_DEFLATED
            member.external_attr = external_attr
            zout.writestr(member, data)
    tmp_path.replace(path)


def _render_docx_body(document: Document, body: str) -> None:
    """Project canonical Markdown semantics to DOCX paragraphs, not raw tokens."""
    for raw in body.split("\n"):
        line = raw.strip()
        if not line or line.startswith("<!--"):
            continue
        if line.startswith("![") and "](" in line:
            continue
        heading = re.match(r"^(#{3,4})\s+(.+)$", line)
        if heading:
            document.add_heading(heading.group(2), level=min(len(heading.group(1)), 4))
            continue
        style = None
        if line.startswith("- "):
            style, line = "List Bullet", line[2:]
        elif line.startswith("> "):
            line = line[2:]
        paragraph = document.add_paragraph(style=style)
        for piece in re.split(r"(\*\*[^*]+\*\*)", line):
            if not piece:
                continue
            if piece.startswith("**") and piece.endswith("**"):
                paragraph.add_run(piece[2:-2]).bold = True
            else:
                paragraph.add_run(piece.replace("`", "").replace("*", ""))


# ---------------------------------------------------------------------------
# Output-root containment (mirror BlueprintProducer)
# ---------------------------------------------------------------------------


def _resolve_output_root(output_root: str | Path) -> Path:
    path = Path(output_root)
    if not path.is_absolute():
        path = REPO_ROOT / path
    # Resolve ``..`` traversal AND follow symlinks before the containment check.
    # ``os.path.normpath`` is LEXICAL — it collapses ``..`` but does NOT follow a
    # symlink, so an in-repo symlink pointing outside the repo would pass a
    # normpath containment check. ``Path.resolve()`` follows the link to its real
    # target, so a symlink escape is caught (mirrors BlueprintProducer discipline,
    # hardened against the symlink case). ``strict=False`` so a not-yet-created
    # output dir still resolves.
    path = path.resolve(strict=False)
    repo_root_resolved = REPO_ROOT.resolve(strict=False)
    try:
        path.relative_to(repo_root_resolved)
    except ValueError as exc:
        raise ValueError(
            f"WorkbookProducer output_root must live under the repo root; got {path}"
        ) from exc
    return path


def _relative_asset_path(path: Path) -> str:
    return path.relative_to(REPO_ROOT).as_posix()


# ---------------------------------------------------------------------------
# Sidecar carrying the audits + DOCX path + DP6 reuse stamp
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class WorkbookSidecar:
    """Run-record sidecar returned alongside the ``ProducedAsset``.

    Carries the artifacts + honesty-gate reports + the DP6 frozen-reuse stamp.
    The canonical MD is the ``ProducedAsset.asset_path``; the DOCX path + audits
    ride here.
    """

    asset: ProducedAsset
    markdown_path: str
    docx_path: str
    numeric_audit: dict
    citation_audit: dict
    segment_coverage: dict[str, bool] | dict[str, object]
    gamma_reuse_justified_by: str | None
    deck_no_regression_eligible: bool
    workbook_brief_receipt: dict[str, object] | None = None
    depth_receipt: dict[str, str] | None = None
    # Q1: visible record of learning objectives whose enrichment overlay did NOT
    # resolve (statement/Bloom degraded to a placeholder). ``None`` when every
    # bound objective resolved (no loss to record).
    lo_overlay_loss: dict[str, object] | None = None


# ---------------------------------------------------------------------------
# WorkbookProducer
# ---------------------------------------------------------------------------


class WorkbookProducer(ModalityProducer):
    """Concrete producer for the ``workbook`` modality (Braid S2)."""

    modality_ref = "workbook"
    status = "ready"

    def __init__(
        self,
        *,
        output_root: str | Path = DEFAULT_WORKBOOK_OUTPUT_ROOT,
        prose_revoicer: ProseRevoicer | None = None,
    ) -> None:
        self._output_root = _resolve_output_root(output_root)
        self._prose_revoicer = prose_revoicer or default_prose_revoicer

    def produce(
        self,
        plan_unit: PlanUnit,
        context: ProductionContext,
        *,
        workbook_title: str | None = None,
        spec: WorkbookSpec,
        segments: Sequence[TranscriptSegment],
        source_text: str,
        citations: Iterable[dict[str, str]] | None = None,
        source_ref_manifest: dict[str, str] | None = None,
        vo_script_text: str = "",
        learning_objectives: Sequence[LearningObjectiveBrief] = (),
        answer_keys: Mapping[str, str] | None = None,
        further_reading: Sequence[FurtherReadingEntry] = (),
        research_entries: Sequence[ResearchEntry] = (),
        research_empty_reason: str | None = None,
        research_omitted_note: str | None = None,
        glossary_articles: Sequence[GlossaryArticleBrief] = (),
        glossary_empty_reason: str | None = None,
        glossary: Any | None = None,
        research_trends: ResearchTrendsBrief | None = None,
        research_supplements: set[str] | None = None,
        diff_files: Iterable[str] | None = None,
        reuse_sha: str = "WORKING",
        pre_work: PreWorkBrief | None = None,
        encounter_mode: Literal["recorded", "live"] = "recorded",
        render_profile: Literal["legacy", "presentation_support"] = "legacy",
        workbook_brief_receipt: dict[str, object] | None = None,
        lo_overlay_loss: dict[str, object] | None = None,
        deep_dive_review: Any | None = None,
    ) -> WorkbookSidecar:
        """Produce the workbook artifacts on the frozen deck inputs.

        Returns a :class:`WorkbookSidecar` (carrying the canonical
        :class:`ProducedAsset` + the DOCX path + honesty-gate reports + the DP6
        reuse stamp). Pure of network I/O; deterministic.
        """
        # id-uniqueness at the binding boundary (filed follow-on).
        assert_unique_collateral_ids(spec)
        # G3 exercise fidelity.
        assert_exercise_fidelity(spec)
        # S1 objective-binding fidelity (no orphan / no phantom objective).
        assert_learning_objective_bindings(spec, learning_objectives)

        # Compose the canonical model (one source of truth).
        doc = compose_workbook(
            plan_unit,
            context,
            spec,
            segments,
            prose_revoicer=self._prose_revoicer,
            workbook_title=workbook_title,
            learning_objectives=learning_objectives,
            answer_keys=answer_keys,
            further_reading=further_reading,
            research_entries=research_entries,
            research_empty_reason=research_empty_reason,
            research_omitted_note=research_omitted_note,
            glossary_articles=glossary_articles,
            glossary_empty_reason=glossary_empty_reason,
            glossary=glossary,
            research_trends=research_trends,
            pre_work=pre_work,
            encounter_mode=encounter_mode,
            render_profile=render_profile,
            lo_overlay_loss_note=(
                str(lo_overlay_loss.get("note")) if lo_overlay_loss else None
            ),
            deep_dive_review=deep_dive_review,
        )
        markdown = render_markdown(doc)

        # --- G1: L2 numeric audit FAIL-mode over the workbook body. ---
        # HONESTY BOUNDARY (do NOT over-claim): symbol-only numeric coverage in
        # FAIL-mode (word-form numerals are NOT gated — named gap
        # ``braid-workbook-wordform-numeral-gap``). ``research_supplements``
        # declares numerals sourced from the run's research leg (beyond the base
        # corpus) so they clear without inflating the source text.
        numeric_audit = audit_workbook_numeric_fidelity(
            markdown, source_text, research_supplements=research_supplements
        )

        # --- G2: citation-fidelity wiring + rendered research DOIs (F-2601). ---
        # Every rendered research_entry.source_ref MUST resolve in the G2 manifest
        # (the DOI section is now UNDER the citation gate, not rendered outside
        # it): a corrupt/absent research source_ref FAILS G2. W2 glossary
        # articles are under the same gate (provenance retained; no fabricate).
        # W3 trend/hot-topic source_refs likewise.
        audited_citations = list(citations or [])
        audited_citations.extend({"source_ref": entry.source_ref} for entry in research_entries)
        # 39.1: the term-keyed projection's COVERED articles ride the same G2
        # gate (uncovered terms carry NO citation by construction — J-1). When
        # a projection is supplied it is the single glossary authority (the
        # ``glossary_articles`` param duplicates its covered articles on the
        # _act path — do not double-audit).
        if glossary is not None:
            audited_citations.extend(
                {"source_ref": article.source_ref}
                for entry in glossary.entries
                for article in entry.articles
            )
        else:
            audited_citations.extend(
                {"source_ref": article.source_ref} for article in glossary_articles
            )
        # 37.2b (R7, non-tautological): the deep-dive G2 leg audits the
        # citation ids parsed from the RENDERED markdown (prose side) against
        # the manifest entries built from the contribution's used rows
        # (structured side, joined in _act.py) — prose/structure divergence
        # FAILS G2. NOTE the honest scope: a used row whose source_ref simply
        # rode into the manifest alongside it cannot fail here (both sides
        # share the used-rows origin); what fails is a rendered marker that no
        # used row backs, or a manifest entry the render contradicts.
        if render_profile == "presentation_support":
            audited_citations.extend(
                deep_dive_g2_citation_entries(markdown, deep_dive_review)
            )
        if research_trends is not None:
            for claim in research_trends.trends:
                if claim.confidence != "unusable" and claim.source_ref:
                    audited_citations.append({"source_ref": claim.source_ref})
            for topic in research_trends.hot_topics:
                if topic.confidence == "unusable":
                    continue
                for source_ref in topic.source_refs:
                    audited_citations.append({"source_ref": source_ref})
        citation_audit = audit_citation_fidelity(
            audited_citations,
            source_ref_manifest or {},
        )

        # --- AC-8: dual-coding non-redundancy (workbook narrative ⊋ VO script). ---
        # The gate ALWAYS rides on the deliverable: when no explicit vo_script is
        # supplied, derive it from the transcript segments (the heard narration is
        # the VO). The comparison is over the NARRATIVE BODY ONLY (depth-delta +
        # re-voiced prose) — structural chrome (headings, "Overview",
        # "References", segment ids) must NOT mask a verbatim-VO body.
        effective_vo = vo_script_text or "\n".join(seg.narration_text for seg in segments)
        narrative_body = _narrative_body_text(doc)
        if render_profile == "legacy":
            self._assert_workbook_superset_of_vo(narrative_body, effective_vo)

        # --- AC-5: segment coverage map (100% coverage, zero phantom). ---
        coverage: dict[str, bool] | dict[str, object]
        coverage = (
            self._segment_coverage(markdown, segments)
            if render_profile == "legacy"
            else {
                "status": "not_applicable_pre_deep_dive",
                "claim_fence": (
                    "No transcript-segment completeness claim until Story 37.2a "
                    "supplies traceable read-prose"
                ),
            }
        )
        # The "no segment dropped" guarantee lives in the PRODUCER, not just the
        # test: every manifest segment must appear in the MD (line-level
        # traceable). Raise if any segment is absent from the deliverable.
        dropped = (
            [seg_id for seg_id, present in coverage.items() if not present]
            if render_profile == "legacy"
            else []
        )
        if dropped:
            raise WorkbookFidelityError(
                f"AC-5: {len(dropped)} transcript segment(s) dropped from the "
                f"workbook (no coverage anchor): {dropped!r}"
            )

        # --- Write canonical MD + DOCX deliverable. ---
        stem = f"{plan_unit.unit_id}@{context.lesson_plan_revision}"
        md_path = self._output_root / f"{stem}.md"
        docx_path = self._output_root / f"{stem}.docx"
        md_path.parent.mkdir(parents=True, exist_ok=True)
        md_path.write_text(markdown, encoding="utf-8")
        render_docx_from_model(doc, docx_path)

        asset = ProducedAsset(
            asset_ref=f"workbook-{stem}",
            modality_ref=self.modality_ref,
            source_plan_unit_id=plan_unit.unit_id,
            asset_path=_relative_asset_path(md_path),
            fulfills=stem,
        )

        # --- DP6 reuse stamp: slide-independent => workbook-only diff is reuse-OK. ---
        reuse_justified_by: str | None
        if diff_files is not None and fresh_gamma_required(diff_files):
            reuse_justified_by = None
        else:
            reuse_justified_by = reuse_stamp(reuse_sha)

        return WorkbookSidecar(
            asset=asset,
            markdown_path=_relative_asset_path(md_path),
            docx_path=_relative_asset_path(docx_path),
            numeric_audit=numeric_audit,
            citation_audit=citation_audit,
            segment_coverage=coverage,
            gamma_reuse_justified_by=reuse_justified_by,
            # A frozen-reuse run is, by construction, ineligible to assert deck
            # no-regression (this story makes no deck-regression claim).
            deck_no_regression_eligible=False,
            workbook_brief_receipt=workbook_brief_receipt,
            depth_receipt=(
                {
                    "status": "not_applicable_pre_deep_dive",
                    "claim_fence": (
                        "No read-prose depth claim until Story 37.2a supplies "
                        "traceable Deep Dive prose"
                    ),
                }
                if render_profile == "presentation_support"
                else None
            ),
            lo_overlay_loss=lo_overlay_loss,
        )

    @staticmethod
    def _segment_coverage(markdown: str, segments: Sequence[TranscriptSegment]) -> dict[str, bool]:
        """AC-5: build {segment_id: present} + reject phantom anchors."""
        manifest_ids = {seg.segment_id for seg in segments}
        coverage = {
            seg.segment_id: (_segment_anchor(seg.segment_id) in markdown) for seg in segments
        }
        # Phantom check: every anchor in the MD must be a real manifest id.
        anchor_re = re.compile(rf"<!--\s*{re.escape(SEGMENT_ANCHOR_PREFIX)}([^\s>]+)\s*-->")
        for match in anchor_re.finditer(markdown):
            anchored = match.group(1)
            if anchored not in manifest_ids:
                raise WorkbookFidelityError(
                    f"AC-5: phantom segment anchor {anchored!r} not in the manifest"
                )
        return coverage

    @staticmethod
    def _assert_workbook_superset_of_vo(markdown: str, vo_script_text: str) -> None:
        """AC-8: workbook narrative is a PROPER superset of the VO script."""
        wb_tokens = set(_tokenize(markdown))
        vo_tokens = set(_tokenize(vo_script_text))
        # Strictly more content than the heard narration alone.
        if not (wb_tokens - vo_tokens):
            raise WorkbookFidelityError(
                "AC-8: workbook narrative carries no depth beyond the VO script "
                "(not a proper superset; dual-coding redundancy)"
            )
        # And the VO content is represented (true superset, not a disjoint rewrite).
        if vo_tokens and not (vo_tokens & wb_tokens):
            raise WorkbookFidelityError(
                "AC-8: workbook does not represent the VO script content "
                "(disjoint rewrite, not a superset)"
            )


# The blocks that constitute the workbook's NARRATIVE body (the dual-coding
# deliverable: re-voiced transcript prose + the depth-delta). Everything else
# ("Overview", "Figures", "Exercises", "References") is structural chrome that
# must NOT count toward the AC-8 superset.
_NARRATIVE_BLOCK_HEADINGS: Final[frozenset[str]] = frozenset(
    {"Transcript-narrative", "Depth-delta narrative"}
)


def _narrative_body_text(doc: _ComposedDoc) -> str:
    """Extract the narrative/depth body text for the AC-8 superset comparison.

    Returns ONLY the prose of the narrative blocks, with structural chrome
    stripped: HTML segment anchors, ``####``/``###`` headings, the
    REVOICE-REQUIRED marker, blockquote markers, and the deterministic
    ``Objective:`` / ``Depth deferred off the slide:`` scaffold lines. What
    remains is the re-voiced prose + the deferred-depth narrative — the content
    that must strictly exceed the heard VO.
    """
    lines: list[str] = []
    for _level, heading, body in doc.blocks:
        if heading not in _NARRATIVE_BLOCK_HEADINGS:
            continue
        for raw in body.split("\n"):
            stripped = raw.strip()
            if not stripped:
                continue
            if stripped.startswith("<!--"):  # segment anchors / markers
                continue
            if stripped.startswith("#"):  # #### segment-id / ### section headings
                continue
            if stripped == REVOICE_REQUIRED_MARKER:
                continue
            if stripped.startswith("Objective:") or stripped.startswith(
                "Depth deferred off the slide:"
            ):
                continue
            # Strip a leading blockquote marker but keep the quoted prose.
            if stripped.startswith(">"):
                stripped = stripped[1:].strip()
            lines.append(stripped)
    return "\n".join(lines)


def _tokenize(text: str) -> list[str]:
    return [t for t in re.findall(r"[a-z0-9]+", text.lower()) if len(t) > 2]


__all__ = [
    "DEFAULT_WORKBOOK_OUTPUT_ROOT",
    "HUMAN_REVIEW_SECTION_HEADING",
    "IRENE_REVIEW_MARKER",
    "NO_READING_PATH_HALO_NOTE",
    "REVOICE_REQUIRED_MARKER",
    "DuplicateCollateralIdError",
    "FurtherReadingEntry",
    "LearningObjectiveBrief",
    "ProseRevoicer",
    "ResearchEntry",
    "TranscriptSegment",
    "WorkbookFidelityError",
    "WorkbookProducer",
    "WorkbookScopeError",
    "WorkbookSidecar",
    "assert_exercise_fidelity",
    "assert_learning_objective_bindings",
    "assert_unique_collateral_ids",
    "audit_citation_fidelity",
    "audit_workbook_numeric_fidelity",
    "compose_workbook",
    "default_prose_revoicer",
    "load_transcript_segments",
    "render_docx_from_model",
    "render_markdown",
]
