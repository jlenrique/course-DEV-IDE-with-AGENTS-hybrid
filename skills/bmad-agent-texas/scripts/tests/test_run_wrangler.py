"""Integration + edge-case tests for skills/bmad-agent-texas/scripts/run_wrangler.py.

Tests are hermetic: no network, no real Notion/Box calls. A synthetic
markdown fixture under tests/fixtures/wrangler-golden/ stands in as both
primary and validation content.

The full suite covers the P0/P1 ACs from story 25-1:
  - happy path -> exit 0 + all 6 artifacts + status=complete
  - cross-validation populated when validation-role asset is present
  - sub-floor completeness -> exit 20 + status=blocked
  - unsupported provider -> exit 20 + known_losses + blocking_issues
  - malformed directive -> exit 30
  - idempotent re-run on existing bundle directory
"""

from __future__ import annotations

import importlib.util
import json
import sys
from pathlib import Path

import pytest
import yaml
from marcus.dispatch.contract import DispatchEnvelope, DispatchReceipt

# Load the runner by path — the hyphenated directory prevents plain import.
_THIS_DIR = Path(__file__).resolve().parent
_RUNNER_PATH = _THIS_DIR.parent / "run_wrangler.py"
_FIXTURE_DIR = _THIS_DIR / "fixtures" / "wrangler-golden"


def _load_runner():
    spec = importlib.util.spec_from_file_location(
        "texas_run_wrangler_under_test", _RUNNER_PATH
    )
    assert spec is not None and spec.loader is not None
    mod = importlib.util.module_from_spec(spec)
    sys.modules["texas_run_wrangler_under_test"] = mod
    spec.loader.exec_module(mod)
    return mod


_runner = _load_runner()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _write_directive(tmp_path: Path, body: dict) -> Path:
    path = tmp_path / "directive.yaml"
    path.write_text(yaml.safe_dump(body, sort_keys=False), encoding="utf-8")
    return path


def _assert_all_six_artifacts(bundle_dir: Path) -> dict[str, Path]:
    """Assert every canonical artifact is present and non-empty."""
    expected = {
        "extracted.md": bundle_dir / "extracted.md",
        "metadata.json": bundle_dir / "metadata.json",
        "manifest.json": bundle_dir / "manifest.json",
        "extraction-report.yaml": bundle_dir / "extraction-report.yaml",
        "ingestion-evidence.md": bundle_dir / "ingestion-evidence.md",
        "result.yaml": bundle_dir / "result.yaml",
    }
    for name, path in expected.items():
        assert path.is_file(), f"Missing artifact: {name}"
        assert path.stat().st_size > 0, f"Empty artifact: {name}"
    return expected


# ---------------------------------------------------------------------------
# AC-2 / AC-7.a — Happy-path end-to-end
# ---------------------------------------------------------------------------


def test_happy_path_primary_only(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-HAPPY-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "APC Module 1 primary",
                    "expected_min_words": 200,
                }
            ],
        },
    )

    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_COMPLETE

    artifacts = _assert_all_six_artifacts(bundle)
    envelope = yaml.safe_load(artifacts["result.yaml"].read_text(encoding="utf-8"))
    assert envelope["status"] == "complete"
    assert envelope["run_id"] == "TEST-HAPPY-001"
    assert len(envelope["materials"]) == 1
    assert envelope["materials"][0]["quality_tier"] == 1  # FULL_FIDELITY
    assert envelope["cross_validation"] == []
    assert envelope["blocking_issues"] == []

    dispatch_contract = envelope["dispatch_contract"]
    validated_envelope = DispatchEnvelope.model_validate(dispatch_contract["envelope"])
    validated_receipt = DispatchReceipt.model_validate(dispatch_contract["receipt"])
    assert validated_envelope.dispatch_kind.value == "texas_retrieval"
    assert validated_receipt.outcome.value == "complete"


# ---------------------------------------------------------------------------
# AC-3 — extraction-report.yaml schema compliance
# ---------------------------------------------------------------------------


def test_extraction_report_schema_compliance(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-SCHEMA-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "APC primary",
                    "expected_min_words": 200,
                }
            ],
        },
    )

    _runner.main(["--directive", str(directive), "--bundle-dir", str(bundle)])

    report = yaml.safe_load((bundle / "extraction-report.yaml").read_text(encoding="utf-8"))

    # Top-level required fields per extraction-report-schema.md v1.0
    for field in (
        "schema_version",
        "run_id",
        "generated_at",
        "overall_status",
        "validator_version",
        "sources",
        "cross_validation",
        "evidence_summary",
    ):
        assert field in report, f"Missing top-level field: {field}"

    assert report["schema_version"] == "1.0"
    assert report["overall_status"] in ("complete", "complete_with_warnings", "blocked")

    # Source entry required fields
    source = report["sources"][0]
    for field in (
        "ref_id",
        "provider",
        "locator",
        "role",
        "tier",
        "tier_value",
        "passed",
        "counts",
        "structural_fidelity",
        "completeness_ratio",
        "extractor_used",
        "fetched_at",
        "evidence",
        "known_losses",
        "recommendations",
    ):
        assert field in source, f"Missing source field: {field}"

    assert source["tier"] in (
        "FULL_FIDELITY",
        "ADEQUATE_WITH_GAPS",
        "DEGRADED",
        "FAILED",
    )
    assert source["tier_value"] in (1, 2, 3, 4)
    for count_field in ("words", "lines", "headings", "expected_min_words"):
        assert count_field in source["counts"]
        assert isinstance(source["counts"][count_field], int)


# ---------------------------------------------------------------------------
# AC-4 — 30-line-stub tripwire preserved
# ---------------------------------------------------------------------------


def test_sub_floor_extraction_blocks_run(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-THIN-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "thin.md"),
                    "role": "primary",
                    "description": "Deliberately thin source",
                    "expected_min_words": 4800,  # 24-page-PDF equivalent
                }
            ],
        },
    )

    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_BLOCKED

    envelope = yaml.safe_load((bundle / "result.yaml").read_text(encoding="utf-8"))
    assert envelope["status"] == "blocked"
    assert envelope["blocking_issues"], "blocked status requires at least one blocking issue"
    assert envelope["materials"][0]["quality_tier"] in (3, 4)  # DEGRADED or FAILED
    # The tripwire should name the insufficient content explicitly
    issue = envelope["blocking_issues"][0]
    assert issue["reason"] == "insufficient_content"


# ---------------------------------------------------------------------------
# AC-6 — Cross-validation populates when validation-role asset is present
# ---------------------------------------------------------------------------


def test_cross_validation_populates_when_validation_asset_present(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-XVAL-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "APC Module 1 primary",
                    "expected_min_words": 200,
                },
                {
                    "ref_id": "src-002",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "validation.md"),
                    "role": "validation",
                    "description": "APC Part 1 reference",
                    "coverage_scope": "part_1_only",
                    "expected_min_words": 50,
                },
            ],
        },
    )

    _runner.main(["--directive", str(directive), "--bundle-dir", str(bundle)])

    envelope = yaml.safe_load((bundle / "result.yaml").read_text(encoding="utf-8"))
    assert envelope["cross_validation"], "cross_validation should not be empty"
    entry = envelope["cross_validation"][0]
    assert entry["primary_ref_id"] == "src-001"
    assert entry["asset_ref_id"] == "src-002"
    assert entry["sections_matched"] >= 1
    assert 0.0 <= entry["key_terms_coverage"] <= 1.0
    assert entry["verdict"]


def test_cross_validation_empty_list_when_no_validation_assets(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-NOXVAL-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "primary only",
                    "expected_min_words": 200,
                }
            ],
        },
    )

    _runner.main(["--directive", str(directive), "--bundle-dir", str(bundle)])
    envelope = yaml.safe_load((bundle / "result.yaml").read_text(encoding="utf-8"))
    # Contract: empty list, NOT null/missing.
    assert envelope["cross_validation"] == []


# ---------------------------------------------------------------------------
# AC-5 — Unsupported provider surfaces as blocked with known_losses
# ---------------------------------------------------------------------------


def test_unsupported_provider_rejected_at_directive_load(
    capsys: pytest.CaptureFixture[str], tmp_path: Path
) -> None:
    """Unknown providers are rejected during directive validation (exit 30)
    so an operator sees the typo immediately instead of burning a fetch
    round on an unreachable source. The runner's remaining fetch-time
    unsupported_provider branch is defensive coverage."""
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-UNSUP-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "carrier_pigeon",  # not a real provider
                    "locator": "not-applicable",
                    "role": "primary",
                    "description": "unsupported",
                }
            ],
        },
    )

    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_DIRECTIVE_OR_IO_ERROR
    err = capsys.readouterr().err
    assert "carrier_pigeon" in err
    assert "provider" in err.lower()


def test_docx_provider_is_accepted_at_directive_load(tmp_path: Path) -> None:
    """`docx` provider should route through local DOCX extraction, not reject."""
    doc = pytest.importorskip("docx")
    bundle = tmp_path / "bundle"
    docx_fixture = tmp_path / "source.docx"
    document = doc.Document()
    document.add_heading("Docx Source", level=1)
    document.add_paragraph("This is a synthetic DOCX source for dispatch acceptance.")
    document.save(docx_fixture)
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-DOCX-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "docx",
                    "locator": str(docx_fixture),
                    "role": "primary",
                    "description": "docx dispatch acceptance",
                    "expected_min_words": 1,
                }
            ],
        },
    )

    exit_code = _runner.main(["--directive", str(directive), "--bundle-dir", str(bundle)])
    assert exit_code in {_runner.EXIT_COMPLETE, _runner.EXIT_COMPLETE_WITH_WARNINGS}

    report = yaml.safe_load((bundle / "extraction-report.yaml").read_text(encoding="utf-8"))
    assert report["sources"][0]["extractor_used"] == "python-docx"


def test_docx_provider_rejects_non_docx_locator(tmp_path: Path) -> None:
    txt_path = tmp_path / "source.txt"
    txt_path.write_text("plain text", encoding="utf-8")

    with pytest.raises(ValueError, match=r"requires a \.docx locator"):
        _runner._fetch_source(
            {
                "provider": "docx",
                "locator": str(txt_path),
                "role": "primary",
                "ref_id": "src-001",
            }
        )


# ---------------------------------------------------------------------------
# AC-7 — Malformed directive exits 30
# ---------------------------------------------------------------------------


def test_malformed_directive_exits_directive_error(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    bad = tmp_path / "bad-directive.yaml"
    bad.write_text(": not valid yaml :\n  - oops\n", encoding="utf-8")

    exit_code = _runner.main(
        ["--directive", str(bad), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_DIRECTIVE_OR_IO_ERROR


def test_missing_directive_file_exits_directive_error(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    missing = tmp_path / "does-not-exist.yaml"

    exit_code = _runner.main(
        ["--directive", str(missing), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_DIRECTIVE_OR_IO_ERROR


def test_directive_missing_required_field_exits_30(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(tmp_path, {"run_id": "X-NO-SOURCES"})

    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_DIRECTIVE_OR_IO_ERROR


# ---------------------------------------------------------------------------
# Idempotency — re-running overwrites cleanly
# ---------------------------------------------------------------------------


def test_idempotent_rerun_on_existing_bundle(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-IDEMPOTENT-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "idempotent",
                    "expected_min_words": 200,
                }
            ],
        },
    )

    first = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert first == _runner.EXIT_COMPLETE
    first_report_mtime = (bundle / "extraction-report.yaml").stat().st_mtime

    second = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert second == _runner.EXIT_COMPLETE
    second_report_mtime = (bundle / "extraction-report.yaml").stat().st_mtime
    # The file should have been rewritten (timestamps move, content is fresh).
    assert second_report_mtime >= first_report_mtime


# ---------------------------------------------------------------------------
# Manifest integrity
# ---------------------------------------------------------------------------


def test_manifest_lists_every_content_artifact_with_sha256(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-MANIFEST-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "manifest check",
                    "expected_min_words": 200,
                }
            ],
        },
    )

    _runner.main(["--directive", str(directive), "--bundle-dir", str(bundle)])

    manifest = json.loads((bundle / "manifest.json").read_text(encoding="utf-8"))
    assert manifest["schema_version"] == "1.0"
    assert manifest["run_id"] == "TEST-MANIFEST-001"
    listed_paths = {artifact["path"] for artifact in manifest["artifacts"]}
    # The four content files (manifest.json itself is excluded by design).
    for required in (
        "extracted.md",
        "metadata.json",
        "extraction-report.yaml",
        "ingestion-evidence.md",
    ):
        assert required in listed_paths, f"Manifest missing: {required}"
    for artifact in manifest["artifacts"]:
        assert len(artifact["sha256"]) == 64  # hex digest length
        assert artifact["size_bytes"] > 0


# ---------------------------------------------------------------------------
# JSON stdout mode
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Post-review regression: tier is re-derived when operator-declared floor applied
# ---------------------------------------------------------------------------


def test_operator_declared_floor_forces_lower_tier(tmp_path: Path) -> None:
    """A directive declaring an expected_min_words floor the extraction cannot
    meet must re-derive the tier downward so the status reflects the operator's
    intent. Without this, a short extraction could pass with FULL_FIDELITY
    because the validator's heuristic guessed a smaller floor."""
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-OVERRIDE-TIER-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "primary w/ aggressive floor",
                    # primary.md has ~340 words. Declaring 4000 expected forces
                    # ratio ~0.085 -> DEGRADED -> blocked.
                    "expected_min_words": 4000,
                }
            ],
        },
    )

    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_BLOCKED

    envelope = yaml.safe_load((bundle / "result.yaml").read_text(encoding="utf-8"))
    assert envelope["status"] == "blocked"
    # Tier must reflect the override, not the validator's own heuristic.
    assert envelope["materials"][0]["quality_tier"] in (3, 4)

    report = yaml.safe_load((bundle / "extraction-report.yaml").read_text(encoding="utf-8"))
    evidence = report["sources"][0]["evidence"]
    assert any("operator-declared floor" in line for line in evidence)
    assert any("Tier re-derived" in line for line in evidence)


# ---------------------------------------------------------------------------
# Post-review regression: tier-2 ADEQUATE_WITH_GAPS path emits complete_with_warnings
# ---------------------------------------------------------------------------


def test_tier_two_emits_complete_with_warnings(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-TIER2-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "primary at ~60% of declared floor",
                    # primary.md ~340 words; 500 floor -> ratio ~0.68 -> ADEQUATE_WITH_GAPS.
                    "expected_min_words": 500,
                }
            ],
        },
    )
    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_COMPLETE_WITH_WARNINGS

    envelope = yaml.safe_load((bundle / "result.yaml").read_text(encoding="utf-8"))
    assert envelope["status"] == "complete_with_warnings"
    assert envelope["materials"][0]["quality_tier"] == 2


# ---------------------------------------------------------------------------
# Post-review regression: role: supplementary is recorded but not extracted
# ---------------------------------------------------------------------------


def test_supplementary_role_recorded_in_provenance_but_not_extracted(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-SUPPL-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "APC primary",
                    "expected_min_words": 200,
                },
                {
                    "ref_id": "src-002",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "validation.md"),
                    "role": "supplementary",
                    "description": "extra context only",
                    "expected_min_words": 50,
                },
            ],
        },
    )
    _runner.main(["--directive", str(directive), "--bundle-dir", str(bundle)])

    # supplementary entry appears in provenance.
    metadata = json.loads((bundle / "metadata.json").read_text(encoding="utf-8"))
    provenance_ref_ids = [p["ref_id"] for p in metadata["provenance"]]
    assert "src-002" in provenance_ref_ids

    # supplementary entry is NOT in extracted.md (title-only section).
    extracted = (bundle / "extracted.md").read_text(encoding="utf-8")
    assert "extra context only" not in extracted

    # supplementary entry is NOT cross-validated.
    envelope = yaml.safe_load((bundle / "result.yaml").read_text(encoding="utf-8"))
    assert envelope["cross_validation"] == []

    # evidence_summary mentions the supplementary source count.
    report = yaml.safe_load((bundle / "extraction-report.yaml").read_text(encoding="utf-8"))
    assert any("supplementary" in line for line in report["evidence_summary"])


# ---------------------------------------------------------------------------
# Post-review regression: directive with no primary source is rejected
# ---------------------------------------------------------------------------


def test_directive_with_no_primary_rejected(
    capsys: pytest.CaptureFixture[str], tmp_path: Path
) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-NO-PRIMARY-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "validation.md"),
                    "role": "validation",
                    "description": "validation only",
                    "expected_min_words": 50,
                },
                {
                    "ref_id": "src-002",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "validation.md"),
                    "role": "supplementary",
                    "description": "supplementary only",
                    "expected_min_words": 50,
                },
            ],
        },
    )
    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_DIRECTIVE_OR_IO_ERROR
    err = capsys.readouterr().err
    assert "primary" in err.lower()


# ---------------------------------------------------------------------------
# Post-review regression: duplicate ref_ids are rejected
# ---------------------------------------------------------------------------


def test_duplicate_ref_ids_rejected(
    capsys: pytest.CaptureFixture[str], tmp_path: Path
) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-DUP-REFID-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "first",
                    "expected_min_words": 200,
                },
                {
                    "ref_id": "src-001",  # duplicate
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "validation.md"),
                    "role": "validation",
                    "description": "second w/ same ref_id",
                    "expected_min_words": 50,
                },
            ],
        },
    )
    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_DIRECTIVE_OR_IO_ERROR
    err = capsys.readouterr().err
    assert "duplicate" in err.lower()


# ---------------------------------------------------------------------------
# Post-review regression: timestamps use Z suffix and are consistent across artifacts
# ---------------------------------------------------------------------------


def test_timestamps_use_z_suffix_and_agree_across_artifacts(tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-TIMESTAMP-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "timestamp check",
                    "expected_min_words": 200,
                }
            ],
        },
    )
    _runner.main(["--directive", str(directive), "--bundle-dir", str(bundle)])

    metadata = json.loads((bundle / "metadata.json").read_text(encoding="utf-8"))
    report = yaml.safe_load((bundle / "extraction-report.yaml").read_text(encoding="utf-8"))
    manifest = json.loads((bundle / "manifest.json").read_text(encoding="utf-8"))

    meta_ts = metadata["generated_at"]
    report_ts = report["generated_at"]
    manifest_ts = manifest["generated_at"]

    # Z suffix, not +00:00.
    for ts in (meta_ts, report_ts, manifest_ts):
        assert ts.endswith("Z"), f"expected Z suffix on timestamp, got {ts!r}"

    # All artifacts for one run agree on the same timestamp string.
    assert meta_ts == report_ts == manifest_ts

    # Per-source fetched_at also uses Z.
    assert metadata["provenance"][0]["fetched_at"].endswith("Z")


# ---------------------------------------------------------------------------
# Original test (restored)
# ---------------------------------------------------------------------------


def test_json_flag_emits_valid_json(capsys: pytest.CaptureFixture[str], tmp_path: Path) -> None:
    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-JSON-001",
            "sources": [
                {
                    "ref_id": "src-001",
                    "provider": "local_file",
                    "locator": str(_FIXTURE_DIR / "primary.md"),
                    "role": "primary",
                    "description": "json mode",
                    "expected_min_words": 200,
                }
            ],
        },
    )

    _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle), "--json"]
    )
    out = capsys.readouterr().out
    payload = json.loads(out)
    assert payload["status"] in ("complete", "complete_with_warnings", "blocked")
    assert payload["run_id"] == "TEST-JSON-001"


# ---------------------------------------------------------------------------
# Story 27-1 — AC-T.5 DOCX integration scenario
# ---------------------------------------------------------------------------


def _build_docx_fixture(path: Path) -> None:
    """Generate a modest synthetic DOCX (~300 words, varied paragraphs)
    suitable for a full run_wrangler integration pass. Each paragraph
    carries distinct content so the validator's line-repetition heuristic
    does not flag the fixture. Generated in-test; no binary in repo."""
    from docx import Document as _DocxDocument

    doc = _DocxDocument()
    doc.add_heading("Integration Fixture: DOCX End-to-End", level=1)
    varied_paragraphs = [
        "This document exercises the full run_wrangler pipeline against a DOCX primary source. "
        "The extractor must open the file via python-docx and produce usable markdown output.",
        "Body-order iteration walks doc.element.body so tables and paragraphs interleave in the "
        "exact sequence authors specified, preserving reading order across extraction.",
        "Heading-style paragraphs become markdown pound-prefix headings. Plain paragraphs "
        "render without prefix, and tables flatten to pipe-separated rows.",
        "The validator inspects the extracted markdown for structural signals such as headings, "
        "paragraph breaks, and content richness before assigning a quality tier.",
        "Fixture diversity matters: if every paragraph repeats, the line-repetition heuristic "
        "in extraction_validator will correctly flag low information density as a known loss.",
    ]
    for para in varied_paragraphs:
        doc.add_paragraph(para)
    doc.add_heading("Subsection Heading", level=2)
    doc.add_paragraph(
        "Additional narrative under the second heading exercises the H2 rendering path and "
        "ensures paragraphs after a heading survive body-order iteration without drift."
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Key"
    tbl.rows[0].cells[1].text = "Value"
    doc.save(str(path))


def test_docx_integration_scenario(tmp_path: Path) -> None:
    """AC-T.5: Directive with a single local_file DOCX source runs end-to-end.

    Asserts: exit code == EXIT_COMPLETE, all 6 artifacts present, extracted.md
    contains DOCX-rendered markdown (heading + pipe-row table), metadata.json
    provenance records kind=local_docx.
    """
    docx_path = tmp_path / "integration_fixture.docx"
    _build_docx_fixture(docx_path)

    bundle = tmp_path / "bundle"
    directive = _write_directive(
        tmp_path,
        {
            "run_id": "TEST-DOCX-INT-001",
            "sources": [
                {
                    "ref_id": "docx-primary",
                    "provider": "local_file",
                    "locator": str(docx_path),
                    "role": "primary",
                    "description": "DOCX integration fixture",
                    "expected_min_words": 150,
                }
            ],
        },
    )

    exit_code = _runner.main(
        ["--directive", str(directive), "--bundle-dir", str(bundle)]
    )
    assert exit_code == _runner.EXIT_COMPLETE

    artifacts = _assert_all_six_artifacts(bundle)

    # extracted.md should contain the rendered heading + table pipe-row,
    # proving python-docx (not read_text_file) produced the output.
    extracted = artifacts["extracted.md"].read_text(encoding="utf-8")
    assert "# Integration Fixture: DOCX End-to-End" in extracted
    assert "## Subsection Heading" in extracted
    assert "| Key | Value |" in extracted
    # Anti-regression: if this text appears, we fell back to read_text_file()
    # and got raw XML/ZIP noise — the bug 27-1 exists to fix.
    assert "PK\x03\x04" not in extracted  # ZIP signature bytes

    # metadata.json provenance records the DOCX-specific extractor label.
    # Note: provenance.kind surfaces the dispatch-level provider ("local_file"),
    # not the SourceRecord.kind. The DOCX-vs-text distinction is carried in
    # provenance.extractor_used, which is populated via _EXTRACTOR_LABELS_BY_KIND
    # when _fetch_source returns a rec with kind "local_docx".
    metadata = json.loads(artifacts["metadata.json"].read_text(encoding="utf-8"))
    provenance = metadata["provenance"]
    assert len(provenance) == 1
    assert provenance[0]["kind"] == "local_file"  # dispatch-level provider
    assert provenance[0]["extractor_used"] == "python-docx"  # extractor-level label
    assert provenance[0]["ref"] == str(docx_path)

    # result.yaml envelope sanity: status=complete, quality_tier=FULL_FIDELITY
    envelope = yaml.safe_load(artifacts["result.yaml"].read_text(encoding="utf-8"))
    assert envelope["status"] == "complete"
    assert envelope["run_id"] == "TEST-DOCX-INT-001"
    assert len(envelope["materials"]) == 1
    assert envelope["materials"][0]["quality_tier"] == 1  # FULL_FIDELITY
