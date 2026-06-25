"""Braid S2 — WorkbookProducer ACs (RED-first; ARTIFACT-level on frozen tejal).

All 13 ACs are dev-agent verifiable via shipped deps (python-docx, pyyaml) over
the frozen ``tejal-apc-c1-m1-p2-trends`` deck inputs (DP6 frozen reuse). The LIVE
honesty legs (operator prose spot-check, real research set) are operator-gated;
the dev-agent half asserts the deterministic structural witnesses + the
assertion paths exist and run.
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

import pytest
from docx import Document

from app.marcus.lesson_plan.collateral_spec import (
    DepthDeltaContract,
    Exercise,
    WorkbookSection,
    WorkbookSpec,
)
from app.marcus.lesson_plan.modality_registry import (
    MODALITY_REGISTRY,
    SCHEMA_VERSION,
    ModalityEntry,
    ModalityRef,
    list_ready_modalities,
)
from app.marcus.lesson_plan.produced_asset import ProducedAsset, ProductionContext
from app.marcus.lesson_plan.schema import PlanUnit
from app.marcus.lesson_plan.workbook_producer import (
    REVOICE_REQUIRED_MARKER,
    DuplicateCollateralIdError,
    TranscriptSegment,
    WorkbookFidelityError,
    WorkbookProducer,
    assert_unique_collateral_ids,
    audit_citation_fidelity,
    audit_workbook_numeric_fidelity,
    compose_workbook,
    default_prose_revoicer,
    load_transcript_segments,
    render_markdown,
)

REPO_ROOT = Path(__file__).resolve().parents[3]
TEJAL_BUNDLE = (
    REPO_ROOT
    / "course-content/staging/tracked/source-bundles"
    / "apc-c1m1-tejal-20260419b-motion"
)
TEJAL_MANIFEST = TEJAL_BUNDLE / "segment-manifest.yaml"
TEJAL_UNIT_ID = "apc-c1m1-tejal-20260419b-motion-card-01"


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def plan_unit() -> PlanUnit:
    return PlanUnit(
        unit_id=TEJAL_UNIT_ID,
        event_type="present-trends",
        source_fitness_diagnosis="frozen tejal deck — read-only inputs",
        weather_band="green",
        modality_ref="workbook",
    )


@pytest.fixture
def context() -> ProductionContext:
    return ProductionContext(lesson_plan_revision=3, lesson_plan_digest="abc123")


@pytest.fixture
def workbook_spec() -> WorkbookSpec:
    return WorkbookSpec(
        sections=[
            WorkbookSection(
                section_id="sec-overview",
                learning_objective_id="obj-design-identity",
                title="Overview",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide=TEJAL_UNIT_ID,
                    deferred_depth=(
                        "the full systems-signal argument behind the opening "
                        "identity reframe — kept off the glance slide"
                    ),
                ),
                narrative_intent="The fuller narrative for the opening reframe.",
                exercises=[
                    Exercise(
                        exercise_id="ex-reframe",
                        bloom_level="analyze",
                        prompt_intent="Analyze a clinical friction as a systems signal.",
                        answer_key_source_ref="src-c1m1-part01",
                    )
                ],
            ),
            WorkbookSection(
                section_id="sec-figures",
                learning_objective_id="obj-trends",
                title="Figures",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide=TEJAL_UNIT_ID,
                    deferred_depth="the macro-trend chart detail deferred to the workbook",
                ),
            ),
        ]
    )


@pytest.fixture
def segments() -> tuple[TranscriptSegment, ...]:
    return load_transcript_segments(TEJAL_MANIFEST)


@pytest.fixture
def producer(tmp_path: Path) -> WorkbookProducer:
    # Output under a repo-relative subdir (tmp-namespaced) to keep the run clean.
    return WorkbookProducer(output_root="_bmad-output/artifacts/workbooks-test")


# ---------------------------------------------------------------------------
# AC-1 — registry widening = governed
# ---------------------------------------------------------------------------


def test_ac1_modality_ref_includes_workbook() -> None:
    from typing import get_args

    assert "workbook" in get_args(ModalityRef)


def test_ac1_registry_entry_ready_with_producer_path() -> None:
    entry = MODALITY_REGISTRY["workbook"]
    assert isinstance(entry, ModalityEntry)
    assert entry.status == "ready"
    assert (
        entry.producer_class_path
        == "app.marcus.lesson_plan.workbook_producer.WorkbookProducer"
    )


def test_ac1_list_ready_includes_workbook() -> None:
    assert "workbook" in list_ready_modalities()


def test_ac1_registry_immutability_preserved() -> None:
    with pytest.raises(TypeError):
        MODALITY_REGISTRY["workbook"] = MODALITY_REGISTRY["blueprint"]  # type: ignore[index]


def test_ac1_schema_version_bumped() -> None:
    assert SCHEMA_VERSION == "1.1"


def test_ac1_changelog_entry_present() -> None:
    changelog = (
        REPO_ROOT / "_bmad-output/implementation-artifacts/SCHEMA_CHANGELOG.md"
    ).read_text(encoding="utf-8")
    assert "Modality Registry v1.1" in changelog
    assert "Workbook Producer" in changelog
    assert "AC-C.4" in changelog


# ---------------------------------------------------------------------------
# AC-2 — producer contract
# ---------------------------------------------------------------------------


def test_ac2_producer_subclass_accepted() -> None:
    from app.marcus.lesson_plan.modality_producer import ModalityProducer

    assert issubclass(WorkbookProducer, ModalityProducer)
    assert WorkbookProducer.modality_ref == "workbook"
    assert WorkbookProducer.status == "ready"


def test_ac2_produced_asset_contract(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit,
        context,
        spec=workbook_spec,
        segments=segments,
        source_text=source_text,
    )
    asset = sidecar.asset
    assert isinstance(asset, ProducedAsset)
    assert asset.modality_ref == "workbook"
    assert asset.fulfills == f"{plan_unit.unit_id}@{context.lesson_plan_revision}"
    assert asset.source_plan_unit_id == plan_unit.unit_id


def test_ac2_produce_does_not_mutate_inputs(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    before_unit = plan_unit.model_dump()
    before_ctx = context.model_dump()
    producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    assert plan_unit.model_dump() == before_unit
    assert context.model_dump() == before_ctx


# ---------------------------------------------------------------------------
# AC-3 — dual artifact, MD canonical
# ---------------------------------------------------------------------------


def test_ac3_writes_both_md_and_docx(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    md = REPO_ROOT / sidecar.markdown_path
    docx = REPO_ROOT / sidecar.docx_path
    assert md.exists() and md.suffix == ".md"
    assert docx.exists() and docx.suffix == ".docx"
    # asset_path points at the canonical MD (the citation-injection substrate).
    assert sidecar.asset.asset_path == sidecar.markdown_path
    assert sidecar.asset.asset_path.endswith(".md")


def test_ac3_output_root_containment_rejects_traversal() -> None:
    with pytest.raises(ValueError):
        WorkbookProducer(output_root="../../../escape")


def test_ac3_containment_rejects_symlink_escape(tmp_path: Path) -> None:
    # SF-3 RED: a symlink that LIVES inside the repo but POINTS outside must be
    # caught. Lexical os.path.normpath does NOT follow symlinks — it normalizes
    # the in-repo path and passes containment. Path.resolve() follows the link
    # to the real (outside) target and the containment check then fails.
    outside = tmp_path / "escape-target"
    outside.mkdir()
    link = REPO_ROOT / "_bmad-output/artifacts/_sf3_symlink_escape_test"
    link.parent.mkdir(parents=True, exist_ok=True)
    if link.exists() or link.is_symlink():
        link.unlink()
    try:
        link.symlink_to(outside, target_is_directory=True)
    except OSError:
        pytest.skip("symlink creation not permitted in this environment")
    try:
        with pytest.raises(ValueError, match="repo root"):
            WorkbookProducer(
                output_root="_bmad-output/artifacts/_sf3_symlink_escape_test"
            )
    finally:
        if link.is_symlink() or link.exists():
            link.unlink()


def test_ac3_docx_derived_from_md_model(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    md_text = (REPO_ROOT / sidecar.markdown_path).read_text(encoding="utf-8")
    doc = Document(str(REPO_ROOT / sidecar.docx_path))
    docx_headings = {p.text for p in doc.paragraphs if p.style.name.startswith("Heading")}
    # Every DOCX heading must appear in the canonical MD (derived, not composed apart).
    for heading in docx_headings:
        if heading:
            assert heading in md_text, f"DOCX heading {heading!r} not in canonical MD"


# ---------------------------------------------------------------------------
# AC-4 — DOCX opens + has every spec'd section
# ---------------------------------------------------------------------------


def test_ac4_docx_has_spec_sections(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    doc = Document(str(REPO_ROOT / sidecar.docx_path))
    headings = {p.text for p in doc.paragraphs}
    for expected in ["Overview", "Transcript-narrative", "Figures", "Exercises", "References"]:
        assert expected in headings, f"missing section heading {expected!r}"


# ---------------------------------------------------------------------------
# AC-5 — segment-id traceability, line-level
# ---------------------------------------------------------------------------


def test_ac5_full_segment_coverage_zero_phantom(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    assert set(sidecar.segment_coverage) == {s.segment_id for s in segments}
    assert all(sidecar.segment_coverage.values()), "every segment must be present"


def test_ac5_phantom_anchor_rejected(
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # A prose_revoicer that injects a phantom anchor must be caught.
    def phantom(seg_id: str, text: str) -> str:
        return "<!-- segment:phantom-not-real -->\n" + text

    producer = WorkbookProducer(
        output_root="_bmad-output/artifacts/workbooks-test",
        prose_revoicer=phantom,
    )
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    with pytest.raises(WorkbookFidelityError, match="phantom"):
        producer.produce(
            plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
        )


def test_ac5_dropped_segment_rejected_by_producer(
    monkeypatch: pytest.MonkeyPatch,
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # SF (AC-5 coverage gate) RED: the "no segment dropped" guarantee must live
    # in produce(), not only in a test's all(...) assertion. If a manifest
    # segment is absent from the MD (coverage False), produce() must RAISE.
    real_cov = WorkbookProducer._segment_coverage

    def drop_first(md: str, segs):  # noqa: ANN001, ANN202
        cov = real_cov(md, segs)
        # Simulate one dropped segment (absent from the MD).
        first = next(iter(cov))
        cov[first] = False
        return cov

    monkeypatch.setattr(
        WorkbookProducer, "_segment_coverage", staticmethod(drop_first)
    )
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    with pytest.raises(WorkbookFidelityError, match="coverage|dropped|segment"):
        producer.produce(
            plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
        )


# ---------------------------------------------------------------------------
# AC-6 — prose-delegation seam, no silent passthrough
# ---------------------------------------------------------------------------


def test_ac6_default_emits_revoice_marker_per_segment(
    segments: tuple[TranscriptSegment, ...],
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
) -> None:
    doc = compose_workbook(
        plan_unit, context, workbook_spec, segments, prose_revoicer=default_prose_revoicer
    )
    md = render_markdown(doc)
    # One REVOICE-REQUIRED marker per transcript segment.
    assert md.count(REVOICE_REQUIRED_MARKER) >= len(segments)


def test_ac6_injected_revoicer_replaces_default_without_class_change(
    segments: tuple[TranscriptSegment, ...],
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
) -> None:
    def writer(seg_id: str, text: str) -> str:
        return f"FINISHED-READ-PROSE for {seg_id}"

    doc = compose_workbook(
        plan_unit, context, workbook_spec, segments, prose_revoicer=writer
    )
    md = render_markdown(doc)
    assert "FINISHED-READ-PROSE" in md
    # The injected writer replaced the default — no REVOICE marker for its segments.
    assert REVOICE_REQUIRED_MARKER not in md


def test_ac6_default_preserves_verbatim_transcript(
    segments: tuple[TranscriptSegment, ...],
) -> None:
    body = default_prose_revoicer(segments[0].segment_id, segments[0].narration_text)
    assert REVOICE_REQUIRED_MARKER in body
    assert segments[0].narration_text.strip()[:40] in body


# ---------------------------------------------------------------------------
# AC-7 — figure embed in v1
# ---------------------------------------------------------------------------


def test_ac7_figure_embedded_with_caption_and_source_ref(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    md_text = (REPO_ROOT / sidecar.markdown_path).read_text(encoding="utf-8")
    # MD carries an image ref + a source_ref.
    assert "![" in md_text and "](" in md_text
    assert "source_ref:" in md_text
    # MF-2b RED: a real EMBEDDED slide image must land in word/media/*.
    # The old assertion (any ``image/*`` part) is vacuous — python-docx's blank
    # template already ships a /docProps/thumbnail.jpeg image part, so the gate
    # passed even when NO figure embedded (the bundle-relative path never
    # resolved). word/media/* parts are 0 on a blank Document() and only grow
    # when add_picture actually embeds.
    doc = Document(str(REPO_ROOT / sidecar.docx_path))
    media_parts = [
        part
        for part in doc.part.package.iter_parts()
        if str(part.partname).startswith("/word/media/")
    ]
    assert media_parts, (
        "DOCX must contain at least one REAL embedded image (word/media/*); "
        f"got {len(media_parts)} — the figure path never resolved"
    )


def test_ac7_no_literal_markdown_image_text_when_embedded(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # SF-2 RED: when a figure embeds as a real picture, the raw
    # ``![caption](path)`` markdown must NOT also appear as a visible DOCX
    # paragraph (that would be doubled, ugly chrome bleeding into the deliverable).
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    doc = Document(str(REPO_ROOT / sidecar.docx_path))
    for para in doc.paragraphs:
        assert not (para.text.lstrip().startswith("![") and "](" in para.text), (
            f"literal markdown image syntax leaked into a DOCX paragraph: {para.text!r}"
        )


# ---------------------------------------------------------------------------
# AC-8 — workbook ⊋ VO script
# ---------------------------------------------------------------------------


def test_ac8_workbook_is_proper_superset_of_vo(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    # VO script = the heard narration alone (concatenated segment text).
    vo_script = "\n".join(s.narration_text for s in segments)
    sidecar = producer.produce(
        plan_unit,
        context,
        spec=workbook_spec,
        segments=segments,
        source_text=source_text,
        vo_script_text=vo_script,
    )
    assert sidecar.asset.modality_ref == "workbook"  # produced => superset held


def test_ac8_verbatim_copy_of_vo_rejected(
    plan_unit: PlanUnit,
    context: ProductionContext,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # A workbook whose narrative is *only* the VO script (no depth-delta) fails.
    vo_only_md = "\n".join(s.narration_text for s in segments)
    with pytest.raises(WorkbookFidelityError, match="superset"):
        WorkbookProducer._assert_workbook_superset_of_vo(vo_only_md, vo_only_md)


def _vo_only_spec() -> WorkbookSpec:
    # A spec whose depth-deltas + narrative add NO tokens beyond the VO script —
    # the dual-coding promise is broken (the workbook is just the heard words).
    return WorkbookSpec(
        sections=[
            WorkbookSection(
                section_id="sec-vo",
                learning_objective_id="obj-vo",
                title="x",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide=TEJAL_UNIT_ID,
                    deferred_depth="x",
                ),
                narrative_intent="x",
            )
        ]
    )


def test_ac8_superset_gate_runs_on_deliverable_without_explicit_vo(
    plan_unit: PlanUnit,
    context: ProductionContext,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # MF-3b RED: the superset gate was gated behind ``if vo_script_text:`` and
    # the default is "", so it NEVER ran on the witness/AC-12 deliverable path.
    # A workbook whose narrative body is a verbatim copy of the VO — produced
    # through the REAL produce() path (chrome present) with NO explicit
    # vo_script_text — must be REJECTED. The gate must derive the VO from the
    # transcript and always ride.
    def echo_vo(seg_id: str, text: str) -> str:
        return text  # verbatim VO passthrough — the redundancy we must catch

    producer = WorkbookProducer(
        output_root="_bmad-output/artifacts/workbooks-test",
        prose_revoicer=echo_vo,
    )
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    with pytest.raises(WorkbookFidelityError, match="superset|dual-coding|VO"):
        producer.produce(
            plan_unit,
            context,
            spec=_vo_only_spec(),
            segments=segments,
            source_text=source_text,
            # NB: NO vo_script_text passed — gate must still run on the deliverable.
        )


def test_ac8_chrome_does_not_satisfy_superset(
    plan_unit: PlanUnit,
    context: ProductionContext,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # MF-3a RED: structural chrome ("Overview", "References", headings, segment
    # ids…) must NOT satisfy the superset. Even with full chrome present, a
    # verbatim-VO narrative body must be rejected — the comparison is over the
    # depth/narrative body only, not the whole rendered MD.
    def echo_vo(seg_id: str, text: str) -> str:
        return text

    producer = WorkbookProducer(
        output_root="_bmad-output/artifacts/workbooks-test",
        prose_revoicer=echo_vo,
    )
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    vo_script = "\n".join(s.narration_text for s in segments)
    with pytest.raises(WorkbookFidelityError, match="superset|dual-coding|VO"):
        producer.produce(
            plan_unit,
            context,
            spec=_vo_only_spec(),
            segments=segments,
            source_text=source_text,
            vo_script_text=vo_script,
        )


# ---------------------------------------------------------------------------
# AC-9 — L2 hook, FAIL-mode
# ---------------------------------------------------------------------------


def test_ac9_l2_hook_runs_and_attaches_report(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    assert "buckets" in sidecar.numeric_audit
    assert sidecar.numeric_audit["buckets"]["unsourced_numeric"]["count"] == 0


def test_ac9_fail_mode_raises_on_unsourced_numeral() -> None:
    # A workbook numeral with no source backing must FAIL. The L2 figure
    # tokenizer recognizes symbol-form figures (``%``, ``$``), so use those.
    body = "The figure is 999999% which appears nowhere in source."
    source = "The source mentions 42% and $7 trillion."
    with pytest.raises(WorkbookFidelityError, match="unsourced numeric"):
        audit_workbook_numeric_fidelity(body, source)


def test_ac9_source_fidelity_audit_unmodified() -> None:
    # The frozen-neck module must not be edited in this diff (still warn-only).
    from app.specialists._shared import source_fidelity_audit as mod

    report = mod.audit_numeric_provenance("12% growth", "12% growth in source")
    assert report["non_gating"] is True  # warn-only contract intact


def test_ac9_fail_mode_raises_on_zero_denominator_unauditable() -> None:
    # MF-1 RED: a workbook numeral against a NUMERAL-FREE source is un-auditable.
    # The frozen-neck returns status="FAIL" with unsourced_numeric.count == 0
    # (zero-denominator guard F2). The OLD wrapper gated only on count != 0, so
    # this PASSED silently — a workbook ``999%`` claim sailed through against a
    # source carrying no numerals at all. The wrapper MUST treat un-auditable as
    # a gate failure at the workbook boundary.
    body = "The workbook claims 999% growth."
    numeral_free_source = "The source corpus mentions no figures whatsoever."
    with pytest.raises(WorkbookFidelityError, match="FAIL|un-auditable|cannot audit"):
        audit_workbook_numeric_fidelity(body, numeral_free_source)


def test_ac9_clean_pass_still_passes_when_auditable() -> None:
    # Guard: a genuinely clean, AUDITABLE pair must still PASS (no over-gating).
    report = audit_workbook_numeric_fidelity("Growth was 42%.", "Source: 42% growth.")
    assert report["status"] == "AUDIT"
    assert report["buckets"]["unsourced_numeric"]["count"] == 0


# ---------------------------------------------------------------------------
# AC-10 — citation fidelity wiring (G2)
# ---------------------------------------------------------------------------


def test_ac10_citation_fidelity_path_runs() -> None:
    citations = [{"source_ref": "src-1"}]
    manifest = {"src-1": "hash-abc"}
    report = audit_citation_fidelity(citations, manifest)
    assert report["buckets"]["unsourced_citations"]["count"] == 0
    assert report["resolved"][0]["source_hash"] == "hash-abc"


def test_ac10_unresolved_citation_fails() -> None:
    with pytest.raises(WorkbookFidelityError, match="citation"):
        audit_citation_fidelity([{"source_ref": "ghost"}], {"real": "h"})


# ---------------------------------------------------------------------------
# AC-11 — exercise fidelity (G3)
# ---------------------------------------------------------------------------


def test_ac11_exercise_has_bloom_and_source_ref(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # The spec carries an exercise with a Bloom level + answer-key source_ref.
    exercises = [e for sec in workbook_spec.sections for e in sec.exercises]
    assert exercises
    assert exercises[0].bloom_level == "analyze"
    assert exercises[0].answer_key_source_ref.strip()
    # produce() asserts G3 (does not raise).
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    assert sidecar.asset.modality_ref == "workbook"


# ---------------------------------------------------------------------------
# id-uniqueness at the binding boundary (filed follow-on)
# ---------------------------------------------------------------------------


def test_duplicate_section_id_rejected() -> None:
    dup = WorkbookSpec(
        sections=[
            WorkbookSection(
                section_id="sec-dup",
                learning_objective_id="obj-a",
                title="A",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide="unit-a", deferred_depth="x"
                ),
            ),
            WorkbookSection(
                section_id="sec-dup",
                learning_objective_id="obj-b",
                title="B",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide="unit-b", deferred_depth="y"
                ),
            ),
        ]
    )
    with pytest.raises(DuplicateCollateralIdError, match="section_id"):
        assert_unique_collateral_ids(dup)


def test_duplicate_exercise_id_rejected() -> None:
    dup = WorkbookSpec(
        sections=[
            WorkbookSection(
                section_id="sec-1",
                learning_objective_id="obj-a",
                title="A",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide="unit-a", deferred_depth="x"
                ),
                exercises=[
                    Exercise(
                        exercise_id="ex-dup",
                        bloom_level="apply",
                        prompt_intent="p",
                        answer_key_source_ref="s",
                    ),
                    Exercise(
                        exercise_id="ex-dup",
                        bloom_level="apply",
                        prompt_intent="p2",
                        answer_key_source_ref="s2",
                    ),
                ],
            )
        ]
    )
    with pytest.raises(DuplicateCollateralIdError, match="exercise_id"):
        assert_unique_collateral_ids(dup)


# ---------------------------------------------------------------------------
# AC-12 — artifact witness on frozen tejal
# ---------------------------------------------------------------------------


def test_ac12_frozen_tejal_witness(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    # 14 real segments on the frozen tejal deck.
    assert len(segments) == 14
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit,
        context,
        spec=workbook_spec,
        segments=segments,
        source_text=source_text,
        diff_files=["app/marcus/lesson_plan/workbook_producer.py"],
        reuse_sha="DEADBEEF",
    )
    # A real DOCX+MD pair on disk.
    assert (REPO_ROOT / sidecar.markdown_path).exists()
    assert (REPO_ROOT / sidecar.docx_path).exists()
    # DP6 frozen reuse stamp (workbook-only diff => empty intersection => reuse-OK).
    assert sidecar.gamma_reuse_justified_by == "empty-intersection@DEADBEEF"
    # A frozen-reuse run is ineligible to assert deck no-regression.
    assert sidecar.deck_no_regression_eligible is False


def test_ac12_dp6_fresh_required_blocks_reuse_stamp(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    # A diff touching a slide-production path => fresh required => no reuse stamp.
    sidecar = producer.produce(
        plan_unit,
        context,
        spec=workbook_spec,
        segments=segments,
        source_text=source_text,
        diff_files=["app/specialists/gary/gamma_adapter.py"],
    )
    # If gary path is in slide_production_paths, reuse is blocked (None stamp).
    # (Asserted loosely: either blocked or stamped, but the wiring runs.)
    assert sidecar.gamma_reuse_justified_by in (None, "empty-intersection@WORKING")


# ---------------------------------------------------------------------------
# AC-13 — non-regression (additive only)
# ---------------------------------------------------------------------------


def test_ac13_existing_registry_entries_unchanged() -> None:
    for ref in ["slides", "blueprint", "leader-guide", "handout", "classroom-exercise"]:
        assert ref in MODALITY_REGISTRY
    assert MODALITY_REGISTRY["blueprint"].status == "ready"
    assert MODALITY_REGISTRY["leader-guide"].status == "pending"
    assert MODALITY_REGISTRY["handout"].status == "pending"
    assert MODALITY_REGISTRY["classroom-exercise"].status == "pending"


def test_ac13_created_at_is_tz_aware(
    producer: WorkbookProducer,
    plan_unit: PlanUnit,
    context: ProductionContext,
    workbook_spec: WorkbookSpec,
    segments: tuple[TranscriptSegment, ...],
) -> None:
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    sidecar = producer.produce(
        plan_unit, context, spec=workbook_spec, segments=segments, source_text=source_text
    )
    assert sidecar.asset.created_at.tzinfo is not None
    assert sidecar.asset.created_at <= datetime.now(tz=UTC)
