"""Braid S3 — workbook S0-S7 v1 cut (RED-first; ARTIFACT-level on the real run).

Closes the design's buildable producer gaps (workbook-component-design-2026-06-25
§7): S1 learning-objectives render (gap 1), S2/S3 transcript-of-record vs
re-voiced split (gap 2), answer-key render (gap 4), citations/further-reading
render (gap 5), and the real tejal WorkbookSpec authoring + a REAL DOCX produced
on the completed B1 run 6a103b6c (gap 7). NO MOCKS — real run + corpus data.
"""

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from app.marcus.lesson_plan.collateral_spec import (
    DepthDeltaContract,
    Exercise,
    WorkbookSection,
    WorkbookSpec,
)
from app.marcus.lesson_plan.prework_projection import (
    FRICTION_SCALE,
    PreWorkBrief,
    PreWorkProvenance,
    PromiseProjection,
    PromiseVow,
    SceneBrief,
)
from app.marcus.lesson_plan.produced_asset import ProductionContext
from app.marcus.lesson_plan.schema import PlanUnit
from app.marcus.lesson_plan.workbook_producer import (
    REVOICE_REQUIRED_MARKER,
    FurtherReadingEntry,
    LearningObjectiveBrief,
    ResearchEntry,
    WorkbookFidelityError,
    WorkbookProducer,
    assert_learning_objective_bindings,
    compose_workbook,
    default_prose_revoicer,
    load_transcript_segments,
    render_docx_from_model,
    render_markdown,
)

REPO_ROOT = Path(__file__).resolve().parents[3]
TEJAL_BUNDLE = (
    REPO_ROOT / "course-content/staging/tracked/source-bundles" / "apc-c1m1-tejal-20260419b-motion"
)
TEJAL_MANIFEST = TEJAL_BUNDLE / "segment-manifest.yaml"
TEST_OUTPUT_ROOT = "_bmad-output/artifacts/workbooks-test"


# --------------------------------------------------------------------------- #
# Fixtures (real frozen-bundle data)                                          #
# --------------------------------------------------------------------------- #


@pytest.fixture
def plan_unit() -> PlanUnit:
    return PlanUnit(
        unit_id="apc-c1m1-tejal-20260419b-motion-card-01",
        event_type="present-trends",
        source_fitness_diagnosis="frozen tejal deck — read-only inputs",
        weather_band="green",
        modality_ref="workbook",
    )


@pytest.fixture
def context() -> ProductionContext:
    return ProductionContext(lesson_plan_revision=3, lesson_plan_digest="abc123")


@pytest.fixture
def segments():
    return load_transcript_segments(TEJAL_MANIFEST)


@pytest.fixture
def two_obj_spec() -> WorkbookSpec:
    return WorkbookSpec(
        sections=[
            WorkbookSection(
                section_id="sec-a",
                learning_objective_id="obj-lo2",
                title="Chapter 2",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide="slide-02",
                    deferred_depth="why administrative waste is an innovation target",
                ),
                narrative_intent="The fuller systems-design argument for chapter 2.",
                exercises=[
                    Exercise(
                        exercise_id="ex-a1",
                        bloom_level="analyze",
                        prompt_intent="Analyze the burnout driver.",
                        answer_key_source_ref="src-slide-02",
                    )
                ],
            ),
            WorkbookSection(
                section_id="sec-b",
                learning_objective_id="obj-lo4",
                title="Chapter 3",
                depth_delta=DepthDeltaContract(
                    deferred_from_slide="slide-05",
                    deferred_depth="the root-cause of the leadership gap",
                ),
                narrative_intent="The root-cause analysis for chapter 3.",
            ),
        ]
    )


def _objectives() -> tuple[LearningObjectiveBrief, ...]:
    return (
        LearningObjectiveBrief(
            "obj-lo2", "analyze", "Analyze the macro-economic and structural trends."
        ),
        LearningObjectiveBrief(
            "obj-lo4", "evaluate", "Evaluate systemic failures via root-cause analysis."
        ),
    )


def _section_body(md: str, heading: str) -> str:
    """Return the body text of a ``## {heading}`` section up to the next ``## ``."""
    marker = f"## {heading}\n"
    start = md.index(marker) + len(marker)
    rest = md[start:]
    nxt = rest.find("\n## ")
    return rest if nxt == -1 else rest[:nxt]


def _prework() -> PreWorkBrief:
    scene = SceneBrief(
        status="authored",
        text="A patient transport delay blocks the work.",
        source_refs=("assessment#Q5",),
        known_losses=(),
        marker=None,
        lesson_type="fresh_pain",
        archetype="external_friction",
    )
    promise = PromiseProjection(
        status="authored",
        vows=(PromiseVow(objective_id="obj-lo2", text="I can identify a first move."),),
        known_losses=(),
        marker=None,
    )
    return PreWorkBrief(
        scene=scene,
        friction_scale=FRICTION_SCALE,
        promise=promise,
        provenance=PreWorkProvenance(source_refs=scene.source_refs, known_losses=()),
    )


def test_presentation_support_frontmatter_and_fr17_cut(
    plan_unit, context, two_obj_spec, segments, tmp_path
) -> None:
    doc = compose_workbook(
        plan_unit,
        context,
        two_obj_spec,
        segments,
        prose_revoicer=default_prose_revoicer,
        pre_work=_prework(),
        encounter_mode="live",
        render_profile="presentation_support",
    )
    md = render_markdown(doc)
    headings = [heading for _, heading, _ in doc.blocks]
    # FLIP-1 (40-1 AC 8, CLOSED A-1 inventory): the cover blocks render FIRST
    # on the presentation_support profile — Cover / Contents / About This
    # Workbook prepend ahead of the pre-work sections.
    assert headings[:4] == ["Cover", "Contents", "About This Workbook", "Pre-work"]
    assert headings[4:7] == ["Scene", "Friction Scale", "Promise"]
    assert "Before the live presentation" in md
    assert "patient transport delay" in md
    assert "Transcript-narrative" not in headings
    assert "Transcript of Record" not in headings
    assert "Figures" not in headings
    assert doc.figures == []
    target = tmp_path / "presentation-support.docx"
    render_docx_from_model(doc, target)
    parsed = Document(target)
    paragraphs = [paragraph.text for paragraph in parsed.paragraphs]
    assert not any("**" in text or text.startswith("- ") for text in paragraphs)
    assert any(paragraph.style.name == "List Bullet" for paragraph in parsed.paragraphs)
    assert "Transcript-narrative" not in paragraphs
    assert "Transcript of Record" not in paragraphs
    assert "Figures" not in paragraphs
    with ZipFile(target) as archive:
        names = archive.namelist()
        assert not any(name.startswith("word/media/") for name in names)
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "/image" not in rels


def test_recorded_live_parity_is_exactly_one_label_in_md_and_docx(
    plan_unit, context, two_obj_spec, segments, tmp_path
) -> None:
    docs = {
        mode: compose_workbook(
            plan_unit,
            context,
            two_obj_spec,
            segments,
            prose_revoicer=default_prose_revoicer,
            pre_work=_prework(),
            encounter_mode=mode,
            render_profile="presentation_support",
        )
        for mode in ("recorded", "live")
    }
    md_lines = {mode: render_markdown(doc).splitlines() for mode, doc in docs.items()}
    differences = [
        (left, right)
        for left, right in zip(md_lines["recorded"], md_lines["live"], strict=True)
        if left != right
    ]
    # FLIP-2 (40-1 AC 8, CLOSED A-1 inventory; consciously EXTENDED per A-3):
    # the cover TOC adds two encounter-copy lines — the Before-you-watch group
    # label and the [presentation] divider. Each copy variant is exactly ONE
    # line with EQUAL line counts across modes, so ``zip(strict=True)`` stays
    # valid and the diff set is exactly the three encounter-copy pairs.
    assert differences == [
        (
            "**Before you watch** — ahead of the recorded presentation",
            "**Before you watch** — ahead of the live presentation",
        ),
        (
            "*[The presentation — watch the recording]*",
            "*[The presentation — attend the live session]*",
        ),
        ("Before you watch the recorded presentation", "Before the live presentation"),
    ]
    paragraphs = {}
    for mode, doc in docs.items():
        path = tmp_path / f"{mode}.docx"
        render_docx_from_model(doc, path)
        paragraphs[mode] = [paragraph.text for paragraph in Document(path).paragraphs]
    differences = [
        (left, right)
        for left, right in zip(paragraphs["recorded"], paragraphs["live"], strict=True)
        if left != right
    ]
    assert differences == [
        (
            "Before you watch — ahead of the recorded presentation",
            "Before you watch — ahead of the live presentation",
        ),
        (
            "[The presentation — watch the recording]",
            "[The presentation — attend the live session]",
        ),
        ("Before you watch the recorded presentation", "Before the live presentation"),
    ]


# --------------------------------------------------------------------------- #
# S1 — learning-objectives section (design §7 gap 1)                          #
# --------------------------------------------------------------------------- #


def test_s1_learning_objectives_section_renders(plan_unit, context, two_obj_spec, segments) -> None:
    doc = compose_workbook(
        plan_unit,
        context,
        two_obj_spec,
        segments,
        prose_revoicer=default_prose_revoicer,
        learning_objectives=_objectives(),
    )
    md = render_markdown(doc)
    body = _section_body(md, "Learning Objectives")
    assert "obj-lo2" in body and "obj-lo4" in body
    # Bloom verb + statement surfaced (was a gap — bindings present, never rendered).
    assert "analyze" in body and "evaluate" in body
    assert "Analyze the macro-economic and structural trends." in body
    # Each objective notes the section(s) that serve it.
    assert "sec-a" in body and "sec-b" in body


def test_s1_thin_render_without_objectives_is_nonregressive(
    plan_unit, context, two_obj_spec, segments
) -> None:
    # No objectives supplied -> thin id-only list derived from the bindings.
    doc = compose_workbook(
        plan_unit, context, two_obj_spec, segments, prose_revoicer=default_prose_revoicer
    )
    md = render_markdown(doc)
    body = _section_body(md, "Learning Objectives")
    assert "obj-lo2" in body and "obj-lo4" in body


def test_s1_orphan_objective_rejected(two_obj_spec) -> None:
    orphan = (*_objectives(), LearningObjectiveBrief("obj-ghost", "remember", "x"))
    with pytest.raises(WorkbookFidelityError, match="orphan"):
        assert_learning_objective_bindings(two_obj_spec, orphan)


def test_s1_phantom_binding_rejected(two_obj_spec) -> None:
    # A section binds obj-lo4, but the objectives list omits it -> phantom binding.
    objs = (LearningObjectiveBrief("obj-lo2", "analyze", "only lo2 stated"),)
    with pytest.raises(WorkbookFidelityError, match="phantom"):
        assert_learning_objective_bindings(two_obj_spec, objs)


def test_s1_produce_raises_on_orphan(plan_unit, context, two_obj_spec, segments) -> None:
    producer = WorkbookProducer(output_root=TEST_OUTPUT_ROOT)
    source_text = (TEJAL_BUNDLE / "extracted.md").read_text(encoding="utf-8")
    with pytest.raises(WorkbookFidelityError, match="orphan"):
        producer.produce(
            plan_unit,
            context,
            spec=two_obj_spec,
            segments=segments,
            source_text=source_text,
            learning_objectives=(
                *_objectives(),
                LearningObjectiveBrief("obj-ghost", "remember", "x"),
            ),
        )


# --------------------------------------------------------------------------- #
# S2/S3 — transcript-of-record distinct from re-voiced explainer (gap 2)      #
# --------------------------------------------------------------------------- #


def test_s3_transcript_of_record_distinct_from_revoiced(
    plan_unit, context, two_obj_spec, segments
) -> None:
    doc = compose_workbook(
        plan_unit, context, two_obj_spec, segments, prose_revoicer=default_prose_revoicer
    )
    md = render_markdown(doc)
    revoiced = _section_body(md, "Transcript-narrative")
    record = _section_body(md, "Transcript of Record")
    # Both carry the verbatim narration of a segment...
    verbatim = segments[0].narration_text.strip()[:50]
    assert verbatim in revoiced
    assert verbatim in record
    # ...but the record is the CLEAN literal: no REVOICE-REQUIRED marker (that is
    # the re-voicing seam, S2-only). The split is real, not a rename.
    assert REVOICE_REQUIRED_MARKER in revoiced
    assert REVOICE_REQUIRED_MARKER not in record
    assert revoiced != record


# --------------------------------------------------------------------------- #
# S5b — answer key render (design §7 gap 4)                                   #
# --------------------------------------------------------------------------- #


def test_s5_answer_key_section_renders_worked_answers(
    plan_unit, context, two_obj_spec, segments
) -> None:
    doc = compose_workbook(
        plan_unit,
        context,
        two_obj_spec,
        segments,
        prose_revoicer=default_prose_revoicer,
        answer_keys={"ex-a1": "Correct answer: bureaucratic and administrative bloat."},
    )
    md = render_markdown(doc)
    body = _section_body(md, "Answer Key")
    assert "ex-a1" in body
    assert "bureaucratic and administrative bloat" in body
    assert "src-slide-02" in body  # answer key grounded by its source_ref


def test_s5_answer_key_pending_note_when_absent(plan_unit, context, two_obj_spec, segments) -> None:
    doc = compose_workbook(
        plan_unit, context, two_obj_spec, segments, prose_revoicer=default_prose_revoicer
    )
    md = render_markdown(doc)
    body = _section_body(md, "Answer Key")
    # No fabricated answer — an explicit pending note + the source_ref.
    assert "pending" in body and "src-slide-02" in body


# --------------------------------------------------------------------------- #
# S6 — citations / further-reading render (design §7 gap 5)                   #
# --------------------------------------------------------------------------- #


def test_s6_citations_and_further_reading_render(
    plan_unit, context, two_obj_spec, segments
) -> None:
    further = (
        FurtherReadingEntry(
            "cit-shrank",
            "Shrank WH et al. Waste in the US Health Care System. JAMA (2019)",
            "src-slide-02",
            "https://doi.org/10.1001/jama.2019.13978",
            "seg-03",
        ),
    )
    research = (
        ResearchEntry(
            "rc-1",
            "A real cited study",
            "src-research-1",
            "scite",
            "10.1234/example.doi",
            supports_segment_id="seg-05",
        ),
    )
    doc = compose_workbook(
        plan_unit,
        context,
        two_obj_spec,
        segments,
        prose_revoicer=default_prose_revoicer,
        further_reading=further,
        research_entries=research,
    )
    md = render_markdown(doc)
    body = _section_body(md, "References")
    assert "Shrank" in body
    assert "https://doi.org/10.1001/jama.2019.13978" in body
    # Live-research DOI rendered as https://doi.org/{source_id}.
    assert "https://doi.org/10.1234/example.doi" in body


def test_s6_empty_research_renders_recorded_empty_not_fabricated(
    plan_unit, context, two_obj_spec, segments
) -> None:
    # S7 D4: zero research rows => recorded explicitly-empty WITH reason (mirror
    # S6 D4). The retired "live-research leg deferred" note is GONE; no DOI is
    # fabricated.
    doc = compose_workbook(
        plan_unit, context, two_obj_spec, segments, prose_revoicer=default_prose_revoicer
    )
    md = render_markdown(doc)
    body = _section_body(md, "References")
    assert "recorded" in body.lower() and "explicitly-empty" in body.lower()
    # No fabricated DOI on an empty research set.
    assert "https://doi.org/" not in body


def test_s6_empty_research_honors_explicit_reason(
    plan_unit, context, two_obj_spec, segments
) -> None:
    # The producer threads an explicit recorded-empty reason (D4); it renders.
    reason = "creds-present run recorded zero cited entries; honest-empty"
    doc = compose_workbook(
        plan_unit,
        context,
        two_obj_spec,
        segments,
        prose_revoicer=default_prose_revoicer,
        research_empty_reason=reason,
    )
    body = _section_body(render_markdown(doc), "References")
    assert reason in body


# --------------------------------------------------------------------------- #
# AC-WITNESS — a REAL DOCX produced on the completed B1 run (design §6 / gap 7)#
# --------------------------------------------------------------------------- #


def test_real_tejal_workbook_docx_witness() -> None:
    from scripts.utilities.produce_tejal_workbook import (
        build_tejal_workbook_inputs,
        produce_tejal_workbook,
    )

    inputs = build_tejal_workbook_inputs()
    # Real run substrate: 13 delivered segments, each with an embedded figure.
    assert len(inputs.segments) == 13
    assert any(s.visual_file_abs for s in inputs.segments)
    # >=3 exercises spanning >=3 Bloom levels (design §6 v1 minimum).
    exercises = [e for sec in inputs.spec.sections for e in sec.exercises]
    assert len(exercises) >= 3
    assert len({e.bloom_level for e in exercises}) >= 3

    sidecar = produce_tejal_workbook(output_root=TEST_OUTPUT_ROOT)
    docx_path = REPO_ROOT / sidecar.docx_path
    md_path = REPO_ROOT / sidecar.markdown_path

    # A real DOCX on disk: non-zero bytes, valid OOXML container (python-docx
    # opens + reads paragraphs).
    assert docx_path.exists() and md_path.exists()
    assert docx_path.stat().st_size > 0
    document = Document(str(docx_path))
    paragraph_texts = [p.text for p in document.paragraphs]
    assert paragraph_texts  # readable

    # S0-S7 section headings present.
    for heading in (
        "Overview",
        "Learning Objectives",
        "Transcript-narrative",
        "Transcript of Record",
        "Figures",
        "Exercises",
        "Answer Key",
        "References",
        "Human Review Checkpoint",
    ):
        assert heading in paragraph_texts, f"missing section heading {heading!r}"

    # >=1 REAL embedded figure (word/media/*), not just the blank-template thumb.
    media_parts = [
        part
        for part in document.part.package.iter_parts()
        if str(part.partname).startswith("/word/media/")
    ]
    assert media_parts, "DOCX must embed >=1 real slide figure (word/media/*)"

    # Honesty gates cleared on the real artifact.
    assert sidecar.numeric_audit["buckets"]["unsourced_numeric"]["count"] == 0
    assert sidecar.citation_audit["buckets"]["unsourced_citations"]["count"] == 0
    assert all(sidecar.segment_coverage.values())

    # Citations section actually rendered the corpus bibliography + a DOI link.
    md_text = md_path.read_text(encoding="utf-8")
    assert "Shrank" in md_text
    assert "https://doi.org/" in md_text
